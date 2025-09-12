# cua_service.py
# -*- coding: utf-8 -*-
import os, json, socket, logging, uuid, time
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional

from sqlalchemy import text
from sqlalchemy.engine import Engine

# ==== Tes modules existants (inchangés) ====
import cerfa_vision_pipeline as cerfa
from intersections_parcelle import (
    get_engine as _get_engine,
    load_layer_map as _load_layer_map,
    locate_parcel_feature as _locate_parcel_feature,
    _intersect_one_parcel as _intersect_one_parcel,
    get_insee_from_csv as _get_insee_from_csv,
)

# Rapport texte (optionnel)
try:
    from report_from_json import generate_text_report as _gen_text_report
except Exception:
    _gen_text_report = None

# Carte HTML (optionnel)
try:
    import folium
    _FOLIUM_OK = True
except Exception:
    _FOLIUM_OK = False

# ========================= Config globale ========================= #
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = Path.cwd() / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

MAPPING_JSON_PATH = os.getenv("MAPPING_JSON_PATH", "mapping_layers.json")
COMMUNES_CSV_PATH = os.getenv("COMMUNES_CSV_PATH", "v_commune_2025.csv")

log = logging.getLogger("cua.service")
if not log.handlers:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    )

# DOCX (optionnel)
try:
    from docx import Document
except Exception:
    Document = None

# Génération CUA DOCX (gabarit structuré)
try:
    from cua_docx import build_cumodel as _build_cumodel, generate_docx_llm as _gen_cua_docx
    _CUA_DOCX_OK = True
except Exception as e:
    _CUA_DOCX_OK = False

# ==== CNIG styles (minimaux, étends si besoin) ====
styles_config = {
    "public.b_zonage_plu": {
        "type": "cnig_plu",
        "field": "typezone",
        "zoom_threshold": 18,
        "small": {
            "U":   {"fill": "rgba(230,0,0,0.6)", "stroke": "#343434", "weight": 0.8},
            "AUc": {"fill": "rgba(255,101,101,0.6)", "stroke": "#343434", "weight": 0.8},
            "AUs": {"fill": "rgba(254,204,190,0.6)", "stroke": "#343434", "weight": 0.8},
            "A":   {"fill": "rgba(255,255,0,0.6)", "stroke": "#343434", "weight": 0.8},
            "N":   {"fill": "rgba(86,170,2,0.6)", "stroke": "#343434", "weight": 0.8}
        },
        "large": {
            "U":   {"stroke": "#B00006", "understroke": "#343434", "dash": "5 2", "weight": 7},
            "AUc": {"stroke": "#D40006", "understroke": "#343434", "dash": "5 2", "weight": 7},
            "AUs": {"stroke": "#E88766", "understroke": "#343434", "dash": "5 2", "weight": 7},
            "A":   {"stroke": "#FFF000", "understroke": "#343434", "dash": "5 2", "weight": 7},
            "N":   {"stroke": "#23A600", "understroke": "#343434", "dash": "5 2", "weight": 7}
        },
        "label_field": "libelle"
    },
    "public.b_assiette_de_servitude_d_utilite_publique": {
        "type": "cnig_sup",
        "zoom_threshold": 18,
        "label_field": None
    }
}

# ==== Template HTML autonome (Leaflet + panneau + couches togglables) ====
HTML_TEMPLATE_BBOX = """<!doctype html>
<html lang="fr">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<title>Intersections BBOX — Parcelle</title>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<style>
:root { --panel-bg:#fff; --panel-shadow:0 8px 22px rgba(0,0,0,.15); --border:#e5e7eb; --muted:#6b7280; }
html, body { height:100%; margin:0; font-family: ui-sans-serif,system-ui,-apple-system,"Segoe UI",Roboto,Arial; }
#map { position:absolute; inset:0; z-index:0; }
.panel { position:absolute; top:12px; left:12px; width:420px; max-height: calc(100% - 24px); overflow:auto;
         background:var(--panel-bg); border:1px solid var(--border); border-radius:12px; box-shadow:var(--panel-shadow); padding:10px; z-index:10000; }
.row { display:flex; align-items:center; gap:8px; justify-content:space-between; }
.small { color:var(--muted); font-size:12px; }
.badge { background:#f3f4f6; border:1px solid var(--border); border-radius:999px; padding:0 8px; font-size:12px; }
.chips { display:flex; flex-wrap:wrap; gap:6px; margin-top:6px; }
.chip { background:#eef2ff; color:#1e40af; border-radius:999px; padding:2px 8px; font-size:12px; }
details.layer { border:1px solid var(--border); border-radius:10px; padding:8px; margin-top:8px; }
details.layer summary { cursor:pointer; font-weight:600; }
.legend-dot { width:10px; height:10px; border-radius:999px; border:1px solid rgba(0,0,0,.2); display:inline-block; }
.controls { display:flex; gap:8px; flex-wrap:wrap; margin:8px 0 4px; }
.controls input { padding:6px 8px; border:1px solid var(--border); border-radius:8px; font-size:13px; width:110px;}
.btn { background:#111827; color:#fff; border:none; border-radius:8px; padding:6px 10px; font-size:13px; cursor:pointer; }
.btn.secondary { background:#f3f4f6; color:#111; border:1px solid var(--border); }
.divider { height:1px; background:var(--border); margin:8px 0; }
.sticky { position:sticky; top:0; background:var(--panel-bg); z-index:5; padding-bottom:6px; }
</style>
</head>
<body>
<div id="map"></div>
<div class="panel">
  <div class="sticky">
    <div class="row"><div><strong>Intersections BBOX</strong></div><div class="small">HTML autonome</div></div>
    <div class="small">Parcelle: <strong id="parcelLbl"></strong> — Buffer: <span id="bufTxt"></span> m</div>
    <div class="small">BBOX: <span id="bboxTxt"></span></div>
    <div class="controls">
      <button id="selectAll" class="btn">Tout cocher</button>
      <button id="clearAll"  class="btn secondary">Tout décocher</button>
      <input id="search" placeholder="filtrer par couche..."/>
    </div>
    <div class="divider"></div>
    <div class="small">Couches intersectantes <span id="hitCount" class="badge">0</span></div>
  </div>
  <div id="layers"></div>
</div>

<script type="application/json" id="data-json">{DATA_JSON}</script>

<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<script>
const DATA = JSON.parse(document.getElementById('data-json').textContent);

const map = L.map('map');
L.tileLayer('https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png',{ maxZoom: 22, attribution: '&copy; OSM' }).addTo(map);

const panelEl = document.querySelector('.panel');
L.DomEvent.disableClickPropagation(panelEl);
L.DomEvent.disableScrollPropagation(panelEl);

document.getElementById('parcelLbl').textContent = DATA.parcel.label;
document.getElementById('bufTxt').textContent = DATA.buffer_m.toFixed(0);
document.getElementById('bboxTxt').textContent = DATA.bbox.join(', ');

const [minx,miny,maxx,maxy] = DATA.bbox;
const bboxPoly = L.polygon([[miny,minx],[miny,maxx],[maxy,maxx],[maxy,minx]], { color:'#111', weight:2, dashArray:'6,6', fillOpacity:0 }).addTo(map);

// parcelle
if (DATA.parcel && DATA.parcel.geojson){
  const parcelLayer = L.geoJSON(DATA.parcel.geojson, {
    style: { color:'#0b5', weight:3, fillOpacity:0.12 }
  }).addTo(map);
  parcelLayer.bindPopup(`<strong>Parcelle</strong><br>${DATA.parcel.label}`);
  map.fitBounds(L.featureGroup([bboxPoly, parcelLayer]).getBounds(), { padding:[16,16] });
} else {
  map.fitBounds(bboxPoly.getBounds(), { padding:[16,16] });
}

const layerObjs = new Map();
function hashColor(name){ let h=0; for(let c of name) h=(h*31 + c.charCodeAt(0))>>>0; const r=(h&255),g=((h>>8)&255),b=((h>>16)&255); return `rgb(${r},${g},${b})`; }

function popupHTML(title, props){
  const head = `<div style="font-weight:600;margin-bottom:4px;">${title}</div>`;
  if(!props || Object.keys(props).length===0) return head + '<i>Pas de propriétés</i>';
  return head + Object.entries(props).map(([k,v])=>`<div><span style="color:#6b7280">${k}</span>: ${v ?? '<i>null</i>'}</div>`).join('');
}

function makePLULayer(item){
  const sty = item.styles || {};
  const field = sty.field || 'typezone';
  const Z = sty.zoom_threshold || 18;
  const small = sty.small || {};
  const large = sty.large || {};

  function styleSmall(f){
    const v = (f.properties||{})[field];
    const cat = small[v] || {};
    return {
      color: cat.stroke || '#343434',
      weight: cat.weight ?? 0.8,
      fillColor: cat.fill || '#ccc',
      fillOpacity: 0.6
    };
  }
  function styleUnder(f){
    const v = (f.properties||{})[field];
    const cat = large[v] || {};
    return { color: cat.understroke || '#343434', weight: 2.5, fillOpacity: 0, dashArray: null };
  }
  function styleOver(f){
    const v = (f.properties||{})[field];
    const cat = large[v] || {};
    return { color: cat.stroke || '#000', weight: (cat.weight || 7), fillOpacity: 0, dashArray: cat.dash || '5 2' };
  }

  const base = L.geoJSON(item.geojson, {
    style: styleSmall,
    onEachFeature: (f,l)=>{
      if (sty.label_field && f.properties && f.properties[sty.label_field]){
        l.bindTooltip(String(f.properties[sty.label_field]), {permanent:false, direction:'center'});
      }
      l.bindPopup(popupHTML(`${item.schema}.${item.table}`, f.properties||{}));
    }
  });

  const under = L.geoJSON(item.geojson, { style: styleUnder });
  const over  = L.geoJSON(item.geojson, { style: styleOver });

  const group = L.featureGroup([base, under, over]);

  function refresh(){
    const z = map.getZoom();
    const smallScale = (z < Z);
    base.setStyle(smallScale ? styleSmall : (_)=>({opacity:0, fillOpacity:0}));
    under.setStyle(smallScale ? (_)=>({opacity:0}) : styleUnder);
    over.setStyle(smallScale ? (_)=>({opacity:0}) : styleOver);
  }
  map.on('zoomend', refresh);
  refresh();
  return group;
}

function supCode(props){
  const raw = String((props && props.idass) || '').toLowerCase();
  return raw.split('-')[0] || '';
}
function supLibelle(props){
  const code = supCode(props);
  const typ = (props && (props.typeass || props.type)) || '';
  return (typ ? `${typ} — ` : '') + (code ? code.toUpperCase() : 'SUP');
}
function supColor(code){
  const fam = (code.match(/^[a-z]+/)||[''])[0];
  const palette = {
    'a':'#8B5E3C','ac':'#E67E22','ar':'#8E44AD','as':'#1ABC9C',
    'el':'#2980B9','i':'#C0392B','int':'#7F8C8D','js':'#F1C40F',
    'pm':'#D35400','pt':'#2C3E50','t':'#27AE60'
  };
  return palette[fam] || '#111';
}
function makeSUPLayer(item){
  const Z = (item.styles && item.styles.zoom_threshold) || 18;
  function styleSmall(f){
    const code = supCode(f.properties||{});
    return { color: supColor(code), weight: 1.2, fillOpacity: 0 };
  }
  function styleUnder(f){ return { color:'#343434', weight:2.5, fillOpacity:0 }; }
  function styleOver(f){
    const code = supCode(f.properties||{});
    return { color: supColor(code), weight:6, fillOpacity:0, dashArray:'6 3' };
  }
  const base  = L.geoJSON(item.geojson, {
    style: styleSmall,
    onEachFeature: (f,l)=>{
      const lbl = supLibelle(f.properties||{});
      l.bindTooltip(lbl, {sticky:true, direction:'top'});
      l.bindPopup(popupHTML(item.name || `${item.schema}.${item.table}`, Object.assign({}, f.properties, {__libelle: lbl})));
    }
  });
  const under = L.geoJSON(item.geojson, { style: styleUnder });
  const over  = L.geoJSON(item.geojson, { style: styleOver });
  const group = L.featureGroup([base, under, over]);
  function refresh(){
    const smallScale = (map.getZoom() < Z);
    base.setStyle(smallScale ? styleSmall : (_)=>({opacity:0, fillOpacity:0}));
    under.setStyle(smallScale ? (_)=>({opacity:0}) : styleUnder);
    over.setStyle(smallScale ? (_)=>({opacity:0}) : styleOver);
  }
  map.on('zoomend', refresh); refresh();
  return group;
}

function addLayerItem(item){
  const wrap = document.getElementById('layers');
  const full = `${item.schema}.${item.table}`;
  const displayName = item.name || full;
  const color = hashColor(full);
  const id = 'chk_' + full.replace(/[^a-z0-9_\\.]/gi,'_');
  const chips = [];
  if (item.values){
    for(const [col,vals] of Object.entries(item.values)){
      vals.slice(0, 12).forEach(v=> chips.push(`<span class="chip" title="${col}">${v}</span>`));
    }
  }
  const det = document.createElement('details');
  det.className = 'layer';
  det.innerHTML = `
    <summary>
      <label style="display:flex; align-items:center; gap:8px;">
        <input type="checkbox" id="${id}">
        <span class="legend-dot" style="background:${color}"></span>
        <span title="${full}">${displayName}</span>
        <span class="badge" title="features">${item.count}</span>
        <span class="small" title="mode d'aperçu" style="margin-left:auto;">${item.preview_mode}</span>
      </label>
    </summary>
    ${chips.length ? `<div class="chips">${chips.join('')}</div>` : ''}
  `;
  wrap.appendChild(det);

  document.getElementById(id).addEventListener('change', (e)=>{
    if (e.target.checked){
      let layer;
      if (item.styles && item.styles.type === 'cnig_plu'){
        layer = makePLULayer(item);
      } else if (item.styles && item.styles.type === 'cnig_sup'){
        layer = makeSUPLayer(item);
      } else {
        const gj = item.geojson;
        layer = L.geoJSON(gj, {
          style: f => ({ color, weight:2, fillOpacity: (f.geometry.type.includes('Polygon')?0.12:0) }),
          pointToLayer: (f, latlng)=> L.circleMarker(latlng, { radius:5, color }),
          onEachFeature: (f, l)=> l.bindPopup(popupHTML(displayName, f.properties || {}))
        });
      }
      layer.addTo(map);
      layerObjs.set(full, layer);
    } else {
      const layer = layerObjs.get(full);
      if (layer){ map.removeLayer(layer); layerObjs.delete(full); }
    }
  });
}

function refreshList(){
  const q = (document.getElementById('search').value || '').toLowerCase();
  const wrap = document.getElementById('layers');
  wrap.innerHTML = '';
  const items = DATA.layers.filter(it => it.count>0 && (
    `${it.schema}.${it.table}`.toLowerCase().includes(q) || 
    (it.name && it.name.toLowerCase().includes(q))
  ));
  document.getElementById('hitCount').textContent = items.length;
  items.forEach(addLayerItem);
}

document.getElementById('search').addEventListener('input', refreshList);
document.getElementById('selectAll').addEventListener('click', ()=>{
  document.querySelectorAll('#layers input[type="checkbox"]').forEach(chk=>{ if(!chk.checked) chk.click(); });
});
document.getElementById('clearAll').addEventListener('click', ()=>{
  document.querySelectorAll('#layers input[type="checkbox"]').forEach(chk=>{ if(chk.checked) chk.click(); });
});

refreshList();
</script>
</body>
</html>"""

# ==== SQL SRID-aware (BBOX reprojetée côté table) ====
COUNT_SQL_BBOX = """
SET LOCAL statement_timeout = '90s';
WITH env AS ( SELECT ST_MakeEnvelope(:minx,:miny,:maxx,:maxy, 4326) AS env4326 ),
env_t AS ( SELECT ST_Transform(env.env4326, :tsrid) AS env_t FROM env ),
cand AS (
  SELECT t.{geom_q} AS g
  FROM {qname} t, env_t e
  WHERE t.{geom_q} IS NOT NULL
    AND t.{geom_q} && e.env_t
)
SELECT COUNT(*)::bigint
FROM cand c
JOIN env_t e ON ST_Intersects(ST_MakeValid(c.g), e.env_t);
"""

VALUES_SQL_BBOX = """
SET LOCAL statement_timeout = '90s';
WITH env AS ( SELECT ST_MakeEnvelope(:minx,:miny,:maxx,:maxy, 4326) AS env4326 ),
env_t AS ( SELECT ST_Transform(env.env4326, :tsrid) AS env_t FROM env ),
cand AS (
  SELECT t.{geom_q} AS g, ({col_sql}::text) AS v
  FROM {qname} t, env_t e
  WHERE t.{geom_q} IS NOT NULL
    AND {col_sql} IS NOT NULL
    AND t.{geom_q} && e.env_t
)
SELECT DISTINCT c.v
FROM cand c
JOIN env_t e ON ST_Intersects(ST_MakeValid(c.g), e.env_t)
LIMIT :lim;
"""

FEATURES_SQL_BBOX = """
SET LOCAL statement_timeout = '120s';
WITH env AS ( SELECT ST_MakeEnvelope(:minx,:miny,:maxx,:maxy, 4326) AS env4326 ),
env_t AS ( SELECT ST_Transform(env.env4326, :tsrid) AS env_t FROM env ),
cand AS (
  SELECT t.{geom_q} AS g, {props_sql} AS props
  FROM {qname} t, env_t e
  WHERE t.{geom_q} IS NOT NULL
    AND t.{geom_q} && e.env_t
),
valid AS ( SELECT ST_MakeValid(c.g) AS g, c.props FROM cand c ),
clip AS (
  SELECT ST_Intersection(v.g, e.env_t) AS g, v.props
  FROM valid v, env_t e
  WHERE ST_Intersects(v.g, e.env_t)
)
SELECT ST_AsGeoJSON(
         CASE WHEN :simp > 0
              THEN ST_SimplifyPreserveTopology(ST_Transform(c.g, 4326), :simp)
              ELSE ST_Transform(c.g, 4326)
          END
       ) AS gj,
       c.props
FROM clip c
LIMIT :maxf;
"""

def qident(*parts): return ".".join(f'"{p}"' for p in parts)

def list_existing_columns(engine: Engine, schema: str, table: str) -> list[str]:
    sql = """
    SELECT column_name FROM information_schema.columns
    WHERE table_schema=:schema AND table_name=:table
    ORDER BY ordinal_position;
    """
    with engine.begin() as con:
        rows = con.execute(text(sql), {"schema": schema, "table": table}).all()
    return [r[0] for r in rows]

def _props_sql(keep_cols):
    if not keep_cols:
        return "NULL::jsonb"
    pairs = ", ".join([f"'{c}', t.\"{c}\"" for c in keep_cols])
    return f"jsonb_build_object({pairs})"

def get_table_srid(eng: Engine, schema: str, table: str, geom_col: str) -> int:
    q = "SELECT Find_SRID(:s,:t,:g)"
    with eng.begin() as con:
        srid = con.execute(text(q), {"s": schema, "t": table, "g": geom_col}).scalar()
    if not srid or int(srid) <= 0:
        q2 = f'SELECT ST_SRID("{geom_col}") FROM "{schema}"."{table}" WHERE "{geom_col}" IS NOT NULL LIMIT 1'
        with eng.begin() as con:
            srid = con.execute(text(q2)).scalar()
    return int(srid or 4326)

def buffer_bbox_from_parcel_geojson(eng: Engine, parcel_geom_geojson: dict, buffer_m: float) -> list[float]:
    q = """
    WITH p AS ( SELECT ST_SetSRID(ST_GeomFromGeoJSON(:gj), 4326) AS g ),
    buf AS ( SELECT ST_Transform(ST_Buffer(ST_Transform(g, 2154), :bufm), 4326) AS g FROM p )
    SELECT ST_XMin(g), ST_YMin(g), ST_XMax(g), ST_YMax(g) FROM buf;
    """
    with eng.begin() as con:
        row = con.execute(text(q), {"gj": json.dumps(parcel_geom_geojson), "bufm": float(buffer_m)}).first()
    if not row: raise RuntimeError("Impossible de calculer la BBOX du buffer.")
    return [float(row[0]), float(row[1]), float(row[2]), float(row[3])]

def build_map_html_bbox(
    *,
    eng: Engine,
    insee: str,
    section: str,
    numero4: str,
    mapping_path: str,
    buffer_m: float = 300.0,
    simplify: float = 0.0,          # 0 = pas de simplification (précision max)
    max_features: int = 1500,
    schema_whitelist: Optional[list[str]] = None,
    out_html: str = "map_bbox_parcelle.html"
):
    # 1) WFS parcelle
    feat = _locate_parcel_feature(insee, section, numero4)
    if not feat or not feat.get("geometry"):
        raise RuntimeError("Parcelle introuvable (WFS)")

    # 2) BBOX (buffer métrique)
    bbox = buffer_bbox_from_parcel_geojson(eng, feat["geometry"], buffer_m=buffer_m)
    minx, miny, maxx, maxy = bbox

    # 3) Chargement mapping
    layers = _load_layer_map(mapping_path)
    if schema_whitelist:
        wl = set(schema_whitelist)
        layers = [l for l in layers if l["schema"] in wl]

    results_html = []
    for lyr in layers:
        schema, table, geom_col = lyr["schema"], lyr["table"], lyr.get("geom_col","geom")
        keep_cols = [c for c in lyr.get("keep", []) if c in set(list_existing_columns(eng, schema, table))]
        qname = qident(schema, table)
        geom_q = f'"{geom_col}"'
        tsrid = get_table_srid(eng, schema, table, geom_col)

        # COUNT
        with eng.begin() as con:
            n = int(con.execute(
                text(COUNT_SQL_BBOX.format(qname=qname, geom_q=geom_q)),
                {"minx":minx,"miny":miny,"maxx":maxx,"maxy":maxy,"tsrid":tsrid}
            ).scalar() or 0)
        if n <= 0:
            continue

        # DISTINCT values
        vals_map = {}
        for col in keep_cols:
            col_sql = f't."{col}"'
            with eng.begin() as con:
                rows = con.execute(
                    text(VALUES_SQL_BBOX.format(qname=qname, geom_q=geom_q, col_sql=col_sql)),
                    {"minx":minx,"miny":miny,"maxx":maxx,"maxy":maxy,"tsrid":tsrid,"lim":100}
                ).all()
            vals_map[col] = [r[0] for r in rows if r[0] is not None]

        # FEATURES (clip à la BBOX, géométrie exacte dans la fenêtre)
        props_sql = _props_sql(keep_cols)
        with eng.begin() as con:
            rows = con.execute(
                text(FEATURES_SQL_BBOX.format(qname=qname, geom_q=geom_q, props_sql=props_sql)),
                {"minx":minx,"miny":miny,"maxx":maxx,"maxy":maxy,"tsrid":tsrid,"simp":float(simplify),"maxf":int(max_features)}
            ).all()

        feats = []
        for gj, props in rows:
            if not gj: continue
            feats.append({
                "type": "Feature", "geometry": json.loads(gj), "properties": props or {}
            })

        table_key = f"{schema}.{table}"
        styles_for_table = styles_config.get(table_key, {})

        results_html.append({
            "schema": schema, "table": table, "name": lyr.get("name", table_key),
            "count": n, "values": vals_map, "preview_mode": "features",
            "geojson": {"type":"FeatureCollection","features":feats},
            "styles": styles_for_table
        })

    # 4) HTML autonome
    data = {
        "parcel": {"label": f"{section} {numero4}", "geojson": feat["geometry"]},
        "buffer_m": float(buffer_m),
        "bbox": [minx, miny, maxx, maxy],
        "layers": results_html
    }
    payload = json.dumps(data, ensure_ascii=False)
    html = HTML_TEMPLATE_BBOX.replace("{DATA_JSON}", payload)
    Path(out_html).write_text(html, encoding="utf-8")
    return out_html


# ========================= Helpers “purs” ========================= #
def _first_nonempty(*vals):
    for v in vals:
        if v is not None and str(v).strip():
            return str(v).strip()
    return None

def _norm_section(s: Optional[str]) -> Optional[str]:
    if not s: return None
    return str(s).strip().upper().replace(" ", "") or None

def _num4(n: Optional[str]) -> Optional[str]:
    if n is None: return None
    n = str(n).strip()
    return n.zfill(4) if n else None

def _dedup_parcels(items: List[Tuple[str,str]]) -> List[Tuple[str,str]]:
    seen, out = set(), []
    for sec, num in items:
        t = (sec, num)
        if sec and num and t not in seen:
            seen.add(t); out.append(t)
    return out

def _dept_hint_from_pages(pages: List[Dict[str, Any]]) -> Optional[str]:
    for pj in pages or []:
        ta = pj.get("terrain_adresse") or {}
        cp = ta.get("code_postal") or (pj.get("coord_demandeur") or {}).get("adresse", {}).get("code_postal")
        if cp:
            cp_str = str(cp).strip()
            if len(cp_str) >= 2 and cp_str[:2].isdigit():
                return cp_str[:2]
        hdr = pj.get("header_cu")
        if isinstance(hdr, dict) and hdr.get("dept"):
            return str(hdr["dept"]).strip().upper()
    return None

def _parse_parcel_refs_inline(parcels_str: str):
    out = []
    for raw in (parcels_str or "").split(","):
        parts = raw.strip().split()
        if len(parts) != 2:
            continue
        sec = parts[0].upper()
        num4 = str(parts[1]).zfill(4)
        out.append((sec, num4))
    return out

def _label_to_slug(label: str) -> str:
    return label.replace(" ", "")

def _write_json(path: Path, obj: dict):
    path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")

def _write_docx(text: str, path: Path) -> bool:
    if not Document:
        log.warning("python-docx non installé — DOCX non généré")
        return False
    doc = Document()
    for para in text.split("\n\n"):
        for line in para.split("\n"):
            doc.add_paragraph(line)
        doc.add_paragraph("")
    doc.save(path)
    return True

def _meta_from_result(result: Dict[str, Any]) -> Dict[str, Any]:
    ctx = result.get("context", {})
    commune = ctx.get("commune") or ""
    insee = ctx.get("insee") or ""
    # première parcelle
    parc = (ctx.get("parcelles") or [{}])[0]
    label = f"{parc.get('section','')} {parc.get('numero','')}".strip()
    return {
        "departement": os.getenv("DEFAULT_DEPARTEMENT", ""),   # optionnel (ou renseigne via couche 'departement')
        "commune": commune,
        "insee": insee,
        "adresse": os.getenv("DEFAULT_ADRESSE", ""),           # si tu as l'adresse ailleurs, remplace
        "references_cadastrales": label,
        "parcelle": label,
        "demandeur": os.getenv("DEFAULT_DEMANDEUR", ""),
        "date_demande": os.getenv("DEFAULT_DATE_DEMANDE", ""),
        "date_arrete": os.getenv("DEFAULT_DATE_ARRETE", ""),
        "numero_arrete": os.getenv("DEFAULT_NUMERO_ARRETE", ""),
        "plu_nom": os.getenv("DEFAULT_PLU_NOM", f"PLU de {commune}"),
        "plu_date_appro": os.getenv("DEFAULT_PLU_DATE", ""),
        # tu peux aussi pousser des articles_* si tu veux forcer des valeurs
    }



# ========================= Blocs métiers ========================= #
def extract_inputs_from_vision(*, pdf_path: str, communes_csv: str) -> Dict[str, Any]:
    pdf = Path(pdf_path).resolve()
    out_json = Path("cerfa_vision_result.json")
    out_dir = Path("cerfa_pages_out")
    res = cerfa.run(pdf, out_json, out_dir)
    if not res.get("success"):
        raise RuntimeError(res.get("error"))

    data = res.get("data") or {}
    pages = res.get("pages", [])

    commune = (data.get("adresse_terrain") or {}).get("commune")
    if not _first_nonempty(commune):
        for pj in pages or []:
            ta = pj.get("terrain_adresse") or {}
            commune = ta.get("commune") or (pj.get("coord_demandeur") or {}).get("adresse", {}).get("commune")
            if _first_nonempty(commune):
                break
    if not _first_nonempty(commune):
        raise RuntimeError("Impossible de déterminer la commune depuis le CERFA")

    log.info("🔎 INSEE via CSV: commune=%s dept=None", commune)
    insee = _get_insee_from_csv(communes_csv, commune, None)
    if not insee:
        dept_hint = _dept_hint_from_pages(pages)
        log.info("🔎 INSEE via CSV: commune=%s dept=%s", commune, dept_hint)
        if dept_hint:
            insee = _get_insee_from_csv(communes_csv, commune, dept_hint)
    if not insee:
        raise RuntimeError(
            f"INSEE introuvable de façon univoque pour la commune '{commune}'. "
            "Corrige le nom de la commune ou fournis le département."
        )

    raw_refs = data.get("references_cadastrales") or []
    parcels = []
    for r in raw_refs:
        sec = _norm_section(r.get("section"))
        num = _num4(r.get("numero"))
        if sec and num:
            parcels.append((sec, num))
    parcels = _dedup_parcels(parcels)
    if not parcels:
        raise RuntimeError("Aucune référence cadastrale valide trouvée")

    return {"insee": insee, "parcels": parcels, "pages": pages}

def get_engine() -> Engine:
    return _get_engine()

def load_layer_map(path: str):
    return _load_layer_map(path)

def intersect_one_parcel_batch(
    *, eng: Engine, layers: List[Dict[str, Any]], insee: str,
    parcels: List[Tuple[str,str]], carve_enclaves: bool,
    enclave_buffer_m: float, values_limit: int
):
    reports, total_hits = [], 0
    for (sec, num4) in parcels:
        log.info("🌐 WFS IGN: %s %s %s", insee, sec, num4)
        feat = _locate_parcel_feature(insee, sec, num4)
        if not feat:
            log.warning("❌ Parcelle non trouvée: %s %s %s", insee, sec, num4)
            reports.append({
                "parcel": {"label": f"{sec} {num4}", "srid": 4326},
                "error": f"Parcelle non trouvée (INSEE {insee}, {sec} {num4})"
            })
            continue
        rpt = _intersect_one_parcel(
            eng=eng, layers=layers, parcel_feature=feat,
            carve_enclaves=carve_enclaves,
            enclave_buffer_m=enclave_buffer_m,
            values_limit=values_limit
        )
        total_hits += int(rpt.get("layers_with_hits", 0) or 0)
        reports.append(rpt)
    return reports, total_hits

def generate_markdown_report_safe(result: Dict[str, Any]) -> str:
    if _gen_text_report:
        txt = _gen_text_report(result).strip()
        return "\n\n".join(p.strip() for p in txt.split("\n\n")) + "\n"
    lines = [
        "# Rapport CUA",
        f"INSEE: {result['context']['insee']}",
        "Parcelles: " + ", ".join([f"{p['section']} {p['numero']}" for p in result['context']['parcelles']]),
        f"Total couches intersectantes: {result['layers_hit_total']}\n"
    ]
    for rep in result.get("reports", []):
        hdr = rep.get("parcel", {}).get("label", "??")
        lines.append(f"## Parcelle {hdr}")
        lines.append(f"- Couches avec hits: {rep.get('layers_with_hits','?')}")
        for r in rep.get("results", []):
            lines.append(f"### {r.get('schema')}.{r.get('table')} (n={r.get('count')})")
            vals = r.get("values") or {}
            for k, v in vals.items():
                if v:
                    preview = ", ".join(map(str, v[:8])) + ("…" if len(v) > 8 else "")
                    lines.append(f"- **{k}**: {preview}")
    return "\n".join(lines) + "\n"



# ========================= Façades “endpoint-ready” ========================= #
def health_check() -> Dict[str, Any]:
    info = {"ok": True, "dns": {}, "db": {}}
    host_env = os.getenv("SUPABASE_HOST")
    region = os.getenv("SUPABASE_REGION", "eu-west-3")
    project_ref = os.getenv("SUPABASE_PROJECT_REF")
    pooler = host_env or (f"aws-0-{region}.pooler.supabase.com" if region else None)
    direct = f"db.{project_ref}.supabase.co" if project_ref else None
    for h in [pooler, direct]:
        if not h:
            continue
        try:
            socket.getaddrinfo(h, None)
            info["dns"][h] = True
        except Exception:
            info["dns"][h] = False
            info["ok"] = False
    try:
        eng = get_engine()
        with eng.begin() as con:
            con.execute(text("select 1"))
        info["db"]["connect"] = True
    except Exception as e:
        info["db"]["connect"] = False
        info["db"]["error"] = str(e)
        info["ok"] = False
    return info

def run_cua_from_pdf(
    *,
    pdf_bytes: bytes,
    filename: str,
    schema_whitelist: Optional[str],
    values_limit: int,
    carve_enclaves: bool,
    enclave_buffer_m: float,
    make_report: bool,
    make_map: bool,
    max_features_per_layer_on_map: int
) -> Dict[str, Any]:
    if not Path(COMMUNES_CSV_PATH).exists():
        raise RuntimeError(f"COMMUNES_CSV_PATH introuvable: {COMMUNES_CSV_PATH}")
    if not Path(MAPPING_JSON_PATH).exists():
        raise RuntimeError(f"MAPPING_JSON_PATH introuvable: {MAPPING_JSON_PATH}")

    up_path = OUTPUT_DIR / f"upload_{filename}"
    up_path.write_bytes(pdf_bytes)

    vision = extract_inputs_from_vision(pdf_path=str(up_path), communes_csv=COMMUNES_CSV_PATH)
    insee = vision["insee"]
    parcels: List[Tuple[str,str]] = vision["parcels"]

    eng = get_engine()
    layers_all = load_layer_map(MAPPING_JSON_PATH)
    if schema_whitelist:
        allowed = set([s.strip() for s in schema_whitelist.split(",") if s.strip()])
        layers_all = [l for l in layers_all if l["schema"] in allowed]

    reports, total_hits = intersect_one_parcel_batch(
        eng=eng,
        layers=layers_all,
        insee=insee,
        parcels=parcels,
        carve_enclaves=bool(carve_enclaves),
        enclave_buffer_m=float(enclave_buffer_m),
        values_limit=int(values_limit)
    )

    result = {
        "success": True,
        "source": "CERFA -> intersections",
        "context": {
            "insee": insee,
            "nb_parcelles": len(parcels),
            "parcelles": [{"section": s, "numero": n} for s, n in parcels],
        },
        "layers_hit_total": total_hits,
        "reports": reports,
    }

    payload: Dict[str, Any] = {"ok": True, "result": result}

    # Rapport
    if make_report:
        md = generate_markdown_report_safe(result)
        md_name = f"rapport_{insee}_{parcels[0][0]}{parcels[0][1]}.md"
        md_path = OUTPUT_DIR / md_name
        md_path.write_text(md, encoding="utf-8")
        payload["report_markdown_url"] = f"/files/{md_name}"

    # Carte
    if make_map:
        try:
            # première parcelle pour l'aperçu
            sec = parcels[0][0]; num4 = parcels[0][1]
            html_name = f"map_bbox_{insee}_{sec}{num4}.html"
            html_path = OUTPUT_DIR / html_name
            build_map_html_bbox(
                eng=eng,
                insee=insee,
                section=sec,
                numero4=num4,
                mapping_path=MAPPING_JSON_PATH,
                buffer_m=300.0,          # ou récupère un paramètre côté API
                simplify=0.0,
                max_features=max_features_per_layer_on_map,
                schema_whitelist=(schema_whitelist.split(",") if isinstance(schema_whitelist, str) and schema_whitelist.strip() else None),
                out_html=str(html_path)
            )
            payload["map_html_url"] = f"/files/{html_name}"
        except Exception as e:
            payload["map_error"] = f"carte BBOX échouée: {e}"

    # --- CUA DOCX gabarit (optionnel) ---
    cua_docx_url = None
    try:
        if _CUA_DOCX_OK:
            meta = _meta_from_result(result)
            cumodel = _build_cumodel(result, meta)  # << utilise le JSON d'intersections
            docx_name = f"CUA_{result['context']['insee']}_{(result['reports'][0].get('parcel') or {}).get('label','PARCELLE').replace(' ','')}.docx"
            docx_path = OUTPUT_DIR / docx_name
            _gen_cua_docx(cumodel, str(docx_path))
            cua_docx_url = f"/files/{docx_name}"
    except Exception as e:
        log.exception(f"Generation CUA.docx échouée: {e}")

    payload["cua_docx_url"] = cua_docx_url

    return payload

def run_cua_direct(
    *,
    parcel: str,
    insee: str = "33234",
    commune: str = "Latresne",
    mapping_path: Optional[str] = None,
    schema_whitelist: Optional[List[str]] = None,
    values_limit: int = 100,
    carve_enclaves: bool = True,
    enclave_buffer_m: float = 120.0,
    make_report: bool = True,
    make_map: bool = True
) -> Dict[str, Any]:
    eng = get_engine()
    mapping_json = mapping_path or str(MAPPING_JSON_PATH)
    layers_all = load_layer_map(mapping_json)
    if schema_whitelist:
        allowed = set(schema_whitelist)
        layers_all = [l for l in layers_all if l["schema"] in allowed]

    refs = _parse_parcel_refs_inline(parcel)
    if not refs:
        raise RuntimeError("Aucune parcelle valide. Format attendu: 'AC 0496' ou 'AC 0496, AC 0497'.")

    log.info("🔎 Direct CUA: %s parcelles, INSEE=%s", len(refs), insee)
    reports, total_hits = [], 0
    for (sec, num4) in refs:
        log.info("🌐 WFS IGN: %s %s %s", insee, sec, num4)
        feat = _locate_parcel_feature(insee, sec, num4)
        if not feat:
            log.warning("❌ Parcelle non trouvée: %s %s %s", insee, sec, num4)
            reports.append({
                "parcel": {"label": f"{sec} {num4}", "srid": 4326},
                "error": f"Parcelle non trouvée (INSEE {insee}, {sec} {num4})"
            })
            continue
        rpt = _intersect_one_parcel(
            eng=eng, layers=layers_all, parcel_feature=feat,
            carve_enclaves=bool(carve_enclaves),
            enclave_buffer_m=float(enclave_buffer_m),
            values_limit=int(values_limit)
        )
        total_hits += int(rpt.get("layers_with_hits", 0) or 0)
        reports.append(rpt)

    result = {
        "success": True,
        "source": "direct",
        "context": {
            "commune": commune,
            "insee": insee,
            "nb_parcelles": len(refs),
            "parcelles": [{"section": s, "numero": n} for s, n in refs],
        },
        "layers_hit_total": total_hits,
        "reports": reports
    }

    # Fichiers de sortie
    insee_tag = str(result["context"]["insee"])
    first_label = (result["reports"][0].get("parcel") or {}).get("label") if result["reports"] else "PARCELLE"
    label_slug = _label_to_slug(first_label or "PARCELLE")

    # JSON complet
    json_path = OUTPUT_DIR / f"result_{insee_tag}_{label_slug}.json"
    _write_json(json_path, result)
    result_json_url = f"/files/{json_path.name}"

    report_md_url = None
    report_docx_url = None
    if make_report:
        try:
            if _gen_text_report:
                report_text = _gen_text_report(result)
            else:
                report_text = generate_markdown_report_safe(result)
            md_path = OUTPUT_DIR / f"rapport_{insee_tag}_{label_slug}.md"
            md_path.write_text(report_text, encoding="utf-8")
            report_md_url = f"/files/{md_path.name}"
            if _write_docx(report_text, OUTPUT_DIR / f"rapport_{insee_tag}_{label_slug}.docx"):
                report_docx_url = f"/files/rapport_{insee_tag}_{label_slug}.docx"
        except Exception as e:
            log.exception(f"Generation rapport échouée: {e}")

    map_html_url = None
    if make_map and refs:
        try:
            sec, num4 = refs[0]
            html_path = OUTPUT_DIR / f"map_bbox_{insee_tag}_{label_slug}.html"
            build_map_html_bbox(
                eng=eng,
                insee=insee_tag,
                section=sec,
                numero4=num4,
                mapping_path=mapping_json,
                buffer_m=float(enclave_buffer_m) if enclave_buffer_m else 300.0,  # réutilise ton param si tu veux
                simplify=0.0,
                max_features=1500,
                schema_whitelist=(schema_whitelist if schema_whitelist else None),
                out_html=str(html_path)
            )
            map_html_url = f"/files/{html_path.name}"
        except Exception as e:
            log.exception(f"Generation carte BBOX échouée: {e}")

    # --- CUA DOCX gabarit (optionnel) ---
    cua_docx_url = None
    try:
        if _CUA_DOCX_OK:
            meta = _meta_from_result(result)
            cumodel = _build_cumodel(result, meta)  # << utilise le JSON d'intersections
            docx_name = f"CUA_{insee_tag}_{label_slug}.docx"
            docx_path = OUTPUT_DIR / docx_name
            _gen_cua_docx(cumodel, str(docx_path))
            cua_docx_url = f"/files/{docx_name}"
    except Exception as e:
        log.exception(f"Generation CUA.docx échouée: {e}")

    return {
        "ok": True,
        "result": result,
        **{k: v for k, v in {
            "result_json_url": result_json_url,
            "report_markdown_url": report_md_url,
            "report_docx_url": report_docx_url,
            "cua_docx_url": cua_docx_url,          # <<<<
            "map_html_url": map_html_url,
        }.items() if v}
    }
