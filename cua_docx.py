# -*- coding: utf-8 -*-
"""
cua_docx.py — Génère un CUA en DOCX selon le gabarit transmis.

Fonctions exposées:
- build_cumodel(data: dict, meta: dict) -> dict
- generate_docx_brut(cumodel: dict, out_path: str) -> str
- generate_docx_llm(cumodel: dict, out_path: str) -> str

Notes:
- On ne parle QUE des couches présentes (donc intersectantes).
- Zonage: on liste TOUTES les zones + % exact (pas de dominante).
- PPR documentaire (ex n_document_pprn_*) va dans Observations, pas dans “Risques opposables”.
- Articles PLU / équipements / taxes ne sont pas dans le JSON → attendus via meta ou laissés “N/D”.
"""

from typing import Dict, Any, List, Optional, Tuple
from collections import defaultdict
from decimal import Decimal, ROUND_HALF_UP
import json
import re
import logging
from datetime import datetime

# DOCX
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Géom/area (optionnel)
try:
    from shapely import wkt as shapely_wkt
    from shapely.ops import transform as shapely_transform
    from pyproj import Transformer
    _GEOM_OK = True
except Exception:
    _GEOM_OK = False

logger = logging.getLogger(__name__)
if not logger.handlers:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# --- LLM (facultatif) : ta lib utilitaire ---
try:
    from llm_utils import call_gpt5_text
    _LLM_OK = True
except Exception:
    _LLM_OK = False
    def call_gpt5_text(*args, **kwargs):
        return {"success": False, "error": "llm_utils not importable", "response": ""}

# --- Règlement PLU (facultatif) : contexte réglementaire ---
try:
    from plu_regulation_context import get_regulation_for_cumodel_zones, join_regulations_for_docx
    _PLU_REG_OK = True
except Exception:
    _PLU_REG_OK = False


# =========================
# Helpers
# =========================
def _norm(s: Optional[str]) -> str:
    return (s or "").strip()

def _normalize_intersections_payload(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Rend compatibles les schémas:
    - Ancien: { results: [...], parcel: {...} }
    - Nouveau: { reports: [ { results: [...] } ], context: {...} }
    Retourne un dict *copie* avec toujours 'results' top-level et 'parcel' complété si possible.
    """
    if data.get("results"):  # déjà ancien schéma
        return data

    d = json.loads(json.dumps(data))  # deep copy
    rep = (d.get("reports") or [{}])[0]
    results = rep.get("results") or []
    d["results"] = results

    # reconstruire 'parcel' minimal si possible
    parcel = d.get("parcel") or {}
    ctx = d.get("context") or {}

    # surface estimée depuis le context si fournie
    if ctx.get("parcel_area_est_m2") and not parcel.get("area_est_m2"):
        parcel["area_est_m2"] = float(ctx["parcel_area_est_m2"])

    # wkt si présent (certaines pipelines le mettent dans context)
    if ctx.get("wkt") and not parcel.get("wkt"):
        parcel["wkt"] = ctx["wkt"]

    d["parcel"] = parcel
    return d

def _round_pct_fraction_to_pct(x: float) -> float:
    """0.50478 -> 50.48"""
    return float(Decimal(x * 100).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))

def _pct_str(pct: Optional[float]) -> str:
    return f"{pct:.2f} %" if pct is not None else "—"

def _find_key_ci(d: Dict[str, Any], candidates: List[str]) -> Optional[str]:
    if not d: return None
    lower = {k.lower(): k for k in d.keys()}
    for c in candidates:
        k = lower.get(c.lower())
        if k: return k
    return None

def _value_attr(layer: Dict[str, Any], names: List[str]) -> List[str]:
    va = layer.get("value_attributes") or {}
    key = _find_key_ci(va, names)
    if not key: return []
    arr = va.get(key) or []
    return [str(x) for x in arr if x is not None]

def _coverage_pairs(layer: Dict[str, Any], names: List[str]) -> List[Tuple[str, float]]:
    cov = layer.get("coverage") or {}
    key = _find_key_ci(cov, names)
    if not key: return []
    out = []
    for row in cov.get(key) or []:
        val = row.get("value")
        pct = float(row.get("parcel_pct") or 0.0)
        out.append((str(val), pct))
    return out

def _is_like(name: str, *needles: str) -> bool:
    n = (name or "").lower()
    return any(k in n for k in needles)

def _classify(table: str) -> str:
    t = (table or "").lower()
    if _is_like(t, "zone_urba", "zonage_plu"):
        return "zonage"
    if _is_like(t, "info_surf", "zones_de_preemptions"):
        return "dpu"
    if _is_like(t, "prescription_surf", "prescriptions_surfaciques"):
        return "prescription_surf"
    if _is_like(t, "prescription_lin", "prescriptions_lineaires", "info_lin"):
        return "prescription_lin"
    if _is_like(t, "sup_"):
        return "sup"
    if _is_like(t, "znieff", "natura"):
        return "env"
    if _is_like(t, "aoc_viticole"):
        return "env"
    if _is_like(t, "irsn_radon"):
        return "infos_generales"
    if _is_like(t, "sismique"):
        return "infos_generales"
    if _is_like(t, "alearg", "rg"):
        return "infos_generales"
    if _is_like(t, "n_document_"):
        return "observations_doc"
    if _is_like(t, "admin_express_departement"):
        return "departement"
    return "autres"

def _parse_date(s: str) -> Optional[str]:
    """
    Retourne une YYYY-MM-DD propre depuis formats: 'YYYYMMDD', 'YYYY-MM-DD', 'YYYY-MM-DD hh:mm:ss'
    """
    if not s: return None
    s = s.strip()
    for fmt in ("%Y%m%d", "%Y-%m-%d", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(s, fmt).strftime("%Y-%m-%d")
        except Exception:
            continue
    # '2021-01-05 00:00:00' (déjà couvert), ou autre → tente split
    try:
        return s.split(" ")[0]
    except Exception:
        return None

def _compute_surface_m2(parcel: Dict[str, Any], results: List[Dict[str, Any]]) -> Optional[float]:
    # 1) Shapely + reprojection (précis)
    try:
        if _GEOM_OK:
            wkt = parcel.get("wkt")
            if wkt:
                geom = shapely_wkt.loads(wkt)
                from shapely.ops import transform as shapely_transform
                from pyproj import Transformer
                tr = Transformer.from_crs(4326, 2154, always_xy=True)
                g_proj = shapely_transform(lambda x, y: tr.transform(x, y), geom)
                area = float(g_proj.area)
                return float(Decimal(area).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))
    except Exception:
        pass

    # 1bis) Fallback: zone déjà fournie par le JSON (nouveau schéma)
    try:
        a = parcel.get("area_est_m2")
        if isinstance(a, (int, float)) and a > 0:
            return float(Decimal(a).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))
    except Exception:
        pass

    # 2) Fallback: prendre un coverage à 100% depuis les layers
    try:
        for lyr in results:
            cov = lyr.get("coverage") or {}
            for arr in cov.values():
                for row in arr or []:
                    if float(row.get("parcel_pct") or 0.0) >= 0.9999:
                        a = float(row.get("area_m2") or 0.0)
                        if a > 0:
                            return float(Decimal(a).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))
    except Exception:
        pass
    return None

def _zonage_from_layers(results: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], Optional[str], Optional[str]]:
    """
    Retourne:
      - zones: [{"libelle","typezone","pct"}...]
      - plu_nom (str)
      - plu_date_appro (YYYY-MM-DD)
    """
    logger.info(f"🔍 Analyse de {len(results)} couches pour extraire le zonage PLU")
    zones_map = defaultdict(lambda: {"libelle": None, "typezone": None, "pct": 0.0})
    plu_date = None
    plu_nom = None

    def ingest(layer: Dict[str, Any]):
        nonlocal plu_date, plu_nom
        table_name = layer.get("table", "N/A")
        logger.info(f"    🔍 Traitement de la couche: {table_name}")
        
        # libellé + % (coverage)
        coverage_pairs = _coverage_pairs(layer, ["libelle", "LIBELLE"])
        logger.info(f"      📊 Paires coverage trouvées: {len(coverage_pairs)}")
        for val, pf in coverage_pairs:
            zones_map[val]["libelle"] = val
            zones_map[val]["pct"] += pf
            logger.info(f"        Zone '{val}' -> pct cumulé: {zones_map[val]['pct']:.4f}")
        
        # typezone (value_attributes)
        vtype = _value_attr(layer, ["typezone", "TYPEZONE"])
        if vtype:
            logger.info(f"      🏷️ Types de zone trouvés: {vtype}")
            # assignation pragmatique: par ordre
            for (key, _), tv in zip(zones_map.items(), vtype):
                zones_map[key]["typezone"] = tv
                logger.info(f"        Zone '{key}' -> typezone: {tv}")
        else:
            logger.info(f"      ⚠️ Aucun typezone trouvé")
        
        # date appro
        for cand in ["datappro", "DATAPPRO"]:
            va = _value_attr(layer, [cand])
            if va:
                d = _parse_date(va[0])
                if d: 
                    plu_date = plu_date or d
                    logger.info(f"      📅 Date appro trouvée: {d}")
        
        # nom PLU (fallback)
        if not plu_nom:
            insee = _value_attr(layer, ["insee", "INSEE"])
            if insee: 
                plu_nom = "PLU de la commune"
                logger.info(f"      🏛️ Nom PLU fallback: {plu_nom}")

    for lyr in results:
        table_name = lyr.get("table") or ""
        if _classify(table_name) == "zonage":
            logger.info(f"  📋 Couche zonage trouvée: {table_name}")
            ingest(lyr)
        else:
            logger.debug(f"  ⏭️ Couche ignorée (pas zonage): {table_name} -> {_classify(table_name)}")

    zones = []
    for z in zones_map.values():
        zones.append({
            "libelle": z["libelle"],
            "typezone": z["typezone"],
            "pct": _round_pct_fraction_to_pct(z["pct"])
        })
    
    zones.sort(key=lambda x: (-x["pct"], x["libelle"] or ""))
    
    logger.info(f"🎯 Zones PLU finales extraites: {len(zones)} zones")
    for i, zone in enumerate(zones):
        logger.info(f"  Zone {i+1}: libelle='{zone['libelle']}', typezone='{zone['typezone']}', pct={zone['pct']}")
    
    logger.info(f"📅 Date appro PLU: {plu_date or 'Non trouvée'}")
    logger.info(f"🏛️ Nom PLU: {plu_nom or 'Non trouvé'}")
    
    return zones, plu_nom, plu_date


# =========================
# 1) Build CUModel
# =========================
def build_cumodel(data: Dict[str, Any], meta: Dict[str, Any]) -> Dict[str, Any]:
    """
    data: JSON d'intersections (dict)
    meta: champs externes (demandeur, dates, références, …)

    Champs meta pris en compte:
      - commune, insee, adresse, references_cadastrales/parcelle
      - date_demande (ou date_depot), date_certificat (ou date_arrete)
      - demandeur, reference_dossier, numero_arrete (optionnel), departement (optionnel)
      - detail_TA, detail_RAP, autres_participations (optionnels)
      - articles_* , disp_* (optionnels si tu as déjà un moteur d'extraction)
      - etat_equipements (optionnel)
    """
    # 🔧 NORMALISATION DU JSON
    data = _normalize_intersections_payload(data)

    results = data.get("results") or []
    parcel = data.get("parcel") or {}

    # Département via layer (fallback meta)
    departement = _norm(meta.get("departement"))
    if not departement:
        for lyr in results:
            if _classify(lyr.get("table") or "") == "departement":
                vals = _value_attr(lyr, ["nom_officiel_en_majuscules", "nom", "NOM"])
                if vals:
                    departement = vals[0].title()
                    break

    # Surface
    surface_m2 = _compute_surface_m2(parcel, results)

    # Zonage + PLU
    zones, plu_nom, plu_date = _zonage_from_layers(results)
    if not plu_nom:
        # défaut propre
        plu_nom = f"PLU de {meta.get('commune') or 'la commune'}"

    # DPU
    dpu_items = []
    for lyr in results:
        if _classify(lyr.get("table") or "") == "dpu":
            agg = defaultdict(float)
            for lib, pf in _coverage_pairs(lyr, ["libelle", "LIBELLE"]):
                agg[lib] += pf
            if agg:
                for k, pf in agg.items():
                    dpu_items.append(f"{k} — {_pct_str(_round_pct_fraction_to_pct(pf))}")
            else:
                labs = _value_attr(lyr, ["libelle", "LIBELLE"]) or ["Droit de préemption urbain"]
                dpu_items.extend(labs)

    dpu_detail = "; ".join(dpu_items) if dpu_items else "Aucun périmètre de DPU intersectant."

    # SUP
    sup_list = []
    for lyr in results:
        if _classify(lyr.get("table") or "") == "sup":
            codes = _value_attr(lyr, ["nom_sup", "id_objet", "code_sup"])
            types = _value_attr(lyr, ["type"])
            for i, c in enumerate(codes or []):
                sup_list.append(f"{c}{f' — {types[i]}' if i < len(types) else ''}")
    liste_SUP = "; ".join(sup_list) if sup_list else "Aucune SUP détectée sur la parcelle."

    # Environnement & risques généraux
    env_detail_items = []
    rga_detail = None
    sismique_detail = None
    ppr_detail = "Aucun plan de prévention des risques (PPR) opposable n’intersecte la parcelle."
    observations = []

    for lyr in results:
        cls = _classify(lyr.get("table") or "")
        if cls == "env":
            # ZNIEFF etc.
            labs = _value_attr(lyr, ["nom_site", "NOM_SITE"])
            cov = _coverage_pairs(lyr, ["nom_site", "NOM_SITE"])
            if cov:
                agg = defaultdict(float)
                for lib, pf in cov:
                    agg[lib] += pf
                for lib, pf in agg.items():
                    env_detail_items.append(f"{lib} — {_pct_str(_round_pct_fraction_to_pct(pf))}")
            elif labs:
                for lib in labs:
                    env_detail_items.append(lib)
        elif cls == "infos_generales":
            if _is_like(lyr.get("table") or "", "alearg", "rg"):
                niv = _value_attr(lyr, ["niveau"])
                alea = _value_attr(lyr, ["alea"])
                if alea:
                    rga_detail = f"{alea[0]}{f' (niveau {niv[0]})' if niv else ''}"
            if _is_like(lyr.get("table") or "", "sismique"):
                sism = _value_attr(lyr, ["Sismicite", "sismicite"])
                if sism:
                    sismique_detail = sism[0]
        elif cls == "observations_doc":
            nom = (_value_attr(lyr, ["nom"]) or _value_attr(lyr, ["nomrisque"]) or _value_attr(lyr, ["nom_site"]))
            site = _value_attr(lyr, ["site_web", "url_fiche", "URL_FICHE"])
            if nom:
                observations.append(f"{nom[0]} — information documentaire (niveau communal).")
            if site:
                observations.append(f"Voir : {site[0]}")

    env_detail = "; ".join(env_detail_items) if env_detail_items else "—"

    # Équipements (faute de données détaillées, on regarde juste une trace 'info_lin')
    etat_equipements = _norm(meta.get("etat_equipements"))
    if not etat_equipements:
        has_aep = any(_classify(lyr.get("table") or "") == "prescription_lin" and
                      ("aep" in " ".join(_value_attr(lyr, ["LIBELLE", "libelle"])).lower())
                      for lyr in results)
        etat_equipements = "Non renseigné (données réseaux non intégrées)." + (" Présence signalée de réseaux AEP." if has_aep else "")

    # Taxes/participations
    detail_TA = _norm(meta.get("detail_TA")) or "Non disponible."
    detail_RAP = _norm(meta.get("detail_RAP")) or "Non disponible."
    autres_participations = _norm(meta.get("autres_participations")) or "—"

    # Dates & identifiants
    date_demande = _norm(meta.get("date_demande") or meta.get("date_depot"))
    date_arrete = _norm(meta.get("date_arrete") or meta.get("date_certificat"))
    numero_arrete = _norm(meta.get("numero_arrete") or meta.get("reference_dossier") or "—")
    references_cadastrales = _norm(meta.get("references_cadastrales") or meta.get("parcelle"))

    # Zonage list string
    if zones:
        zonage_items = []
        for z in zones:
            libelle = z['libelle']
            typezone = z.get('typezone')
            pct = _pct_str(z['pct'])
            
            if typezone:
                item = f"{libelle} ({typezone}) – {pct}"
            else:
                item = f"{libelle} – {pct}"
            
            zonage_items.append(item)
        zonage_list = "; ".join(zonage_items)
    else:
        zonage_list = "—"

    cumodel = {
        "header": {
            "departement": departement or "—",
            "commune": _norm(meta.get("commune")),
            "numero_arrete": numero_arrete
        },
        "vu": {
            "plu_nom": plu_nom,
            "plu_date_appro": plu_date or "—",
            "demandeur": _norm(meta.get("demandeur")),
            "date_demande": date_demande or "—"
        },
        "article2": {
            "commune": _norm(meta.get("commune")),
            "insee": _norm(meta.get("insee")),
            "adresse_ou_localisation": _norm(meta.get("adresse")),
            "references_cadastrales": references_cadastrales,
            "surface_m2": f"{surface_m2:.2f}" if isinstance(surface_m2, (int, float)) else "—",
            "plu_nom": plu_nom,
            "plu_date_appro": plu_date or "—"
        },
        "article3": {
            "zonage_list": zonage_list,
            # champs “articles_* / disp_*” : laissés vides (N/D) sauf si meta les fournit
            "articles_implantation_voies": _norm(meta.get("articles_implantation_voies")),
            "disp_implantation_voies": _norm(meta.get("disp_implantation_voies")),
            "articles_implantation_limites": _norm(meta.get("articles_implantation_limites")),
            "disp_implantation_limites": _norm(meta.get("disp_implantation_limites")),
            "articles_hauteurs": _norm(meta.get("articles_hauteurs")),
            "disp_hauteurs": _norm(meta.get("disp_hauteurs")),
            "articles_emprise": _norm(meta.get("articles_emprise")),
            "disp_emprise": _norm(meta.get("disp_emprise")),
            "articles_stationnement": _norm(meta.get("articles_stationnement")),
            "disp_stationnement": _norm(meta.get("disp_stationnement")),
            "articles_espaces_verts": _norm(meta.get("articles_espaces_verts")),
            "disp_espaces_verts": _norm(meta.get("disp_espaces_verts")),
            "articles_acces_voirie": _norm(meta.get("articles_acces_voirie")),
            "disp_acces_voirie": _norm(meta.get("disp_acces_voirie")),
            "articles_aspect": _norm(meta.get("articles_aspect")),
            "disp_aspect": _norm(meta.get("disp_aspect")),
            "articles_annexes": _norm(meta.get("articles_annexes")),
            "disp_annexes": _norm(meta.get("disp_annexes")),
        },
        "article4": {
            "liste_SUP": liste_SUP
        },
        "article5": {
            "ppr_detail": ppr_detail,
            "rga_detail": rga_detail or "—",
            "sismique_detail": sismique_detail or "—",
            "env_detail": env_detail,
            "autres_infos": "; ".join(observations) if observations else "—"
        },
        "article6": {
            "etat_equipements": etat_equipements
        },
        "article7": {
            "detail_TA": detail_TA,
            "detail_RAP": detail_RAP,
            "autres_participations": autres_participations
        },
        "article8": {
            "dpu_detail": dpu_detail
        },
        "footer": {
            "commune": _norm(meta.get("commune")),
            "date_arrete": date_arrete or "—",
            "maire": _norm(meta.get("maire") or "Le Maire")
        }
    }
    
    # Expose aussi les zones au LLM/réglementation
    cumodel["plu"] = {"zones": zones}
    
    return cumodel


# =========================
# Mise en forme DOCX (gabarit)
# =========================
def _ensure_styles(doc: Document):
    styles = doc.styles
    # Normal
    stn = styles["Normal"]
    stn.font.name = "Calibri"
    stn.font.size = Pt(11)

def _p(doc: Document, txt: str, bold: bool=False, center: bool=False):
    p = doc.add_paragraph(txt)
    if bold:
        for r in p.runs: r.bold = True
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def _kv(doc: Document, label: str, value: str):
    # ligne simple "Label  Value"
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    p.add_run(_norm(value))

def _article3_table(doc: Document, a3: Dict[str, Any]):
    # Table 3 colonnes: Thématique / Référence / Disposition
    t = doc.add_table(rows=1, cols=3)
    hdr = t.rows[0].cells
    hdr[0].text = "Thématique"
    hdr[1].text = "Référence (article[s])"
    hdr[2].text = "Disposition (extrait neutre)"

    rows = [
        ("Implantation par rapport aux voies/alignements",
         a3.get("articles_implantation_voies") or "N/D",
         a3.get("disp_implantation_voies") or "N/D"),
        ("Implantation par rapport aux limites séparatives",
         a3.get("articles_implantation_limites") or "N/D",
         a3.get("disp_implantation_limites") or "N/D"),
        ("Hauteurs / gabarits",
         a3.get("articles_hauteurs") or "N/D",
         a3.get("disp_hauteurs") or "N/D"),
        ("Emprise au sol / CES",
         a3.get("articles_emprise") or "N/D",
         a3.get("disp_emprise") or "N/D"),
        ("Stationnement",
         a3.get("articles_stationnement") or "N/D",
         a3.get("disp_stationnement") or "N/D"),
        ("Espaces verts / plantations",
         a3.get("articles_espaces_verts") or "N/D",
         a3.get("disp_espaces_verts") or "N/D"),
        ("Accès et voirie",
         a3.get("articles_acces_voirie") or "N/D",
         a3.get("disp_acces_voirie") or "N/D"),
        ("Aspect extérieur / matériaux",
         a3.get("articles_aspect") or "N/D",
         a3.get("disp_aspect") or "N/D"),
        ("Annexes et constructions accessoires",
         a3.get("articles_annexes") or "N/D",
         a3.get("disp_annexes") or "N/D"),
    ]
    for (th, refa, disp) in rows:
        r = t.add_row().cells
        r[0].text = th
        r[1].text = refa
        r[2].text = disp


# =========================
# 2) DOCX — BRUT (gabarit)
# =========================
def generate_docx_brut(cumodel: Dict[str, Any], out_path: str) -> str:
    doc = Document()
    _ensure_styles(doc)

    H = cumodel["header"]; Vu = cumodel["vu"]; A2 = cumodel["article2"]
    A3 = cumodel["article3"]; A4 = cumodel["article4"]; A5 = cumodel["article5"]
    A6 = cumodel["article6"]; A7 = cumodel["article7"]; A8 = cumodel["article8"]; F = cumodel["footer"]

    # En-tête République Française
    _p(doc, "République Française", bold=True, center=True)
    _p(doc, f"Département de {H['departement']}", center=True)
    _p(doc, f"Commune de {H['commune']}", center=True)
    doc.add_paragraph()
    _p(doc, f"ARRÊTÉ n° {H['numero_arrete']}", bold=True, center=True)
    _p(doc, "Portant certificat d’urbanisme d’information (CU a)", center=True)
    doc.add_paragraph()

    # Vu
    _p(doc, "Vu", bold=True)
    _p(doc, "— le Code de l’urbanisme, notamment ses articles L.410-1 et R.410-1 et suivants ;")
    _p(doc, f"— le plan local d’urbanisme (PLU) de {H['commune']}, approuvé le {Vu['plu_date_appro']} ;")
    _p(doc, f"— la demande de certificat d’urbanisme d’information déposée par {Vu['demandeur']}, reçue le {Vu['date_demande']} ;")
    doc.add_paragraph()

    # Article 1
    _p(doc, "Article 1 – Objet", bold=True)
    _p(doc, ("Le présent arrêté vaut certificat d’urbanisme d’information (CU a) et indique l’état des "
             "règles d’urbanisme applicables au terrain à la date de sa signature. Il ne vaut ni "
             "autorisation d’occuper ou d’utiliser le sol, ni attestation de faisabilité technique ou "
             "foncière."))
    doc.add_paragraph()

    # Article 2
    _p(doc, "Article 2 – Identification et localisation du terrain", bold=True)
    _kv(doc, "Commune", f"{A2['commune']} ({A2['insee']})")
    _kv(doc, "Adresse / Localisation", A2["adresse_ou_localisation"])
    _kv(doc, "Références cadastrales", A2["references_cadastrales"])
    _kv(doc, "Surface indicative", f"{A2['surface_m2']} m²")
    _kv(doc, "Document d’urbanisme opposable", f"{A2['plu_nom']} (approuvé le {A2['plu_date_appro']})")
    doc.add_paragraph()

    # Article 3
    _p(doc, "Article 3 – Dispositions d’urbanisme applicables (PLU)", bold=True)
    _kv(doc, "Zonage", A3["zonage_list"])
    _p(doc, ("Les occupations et utilisations du sol, ainsi que les règles de constructibilité, sont "
             "définies par le règlement du PLU. Ci-dessous, les thématiques majeures sont rappelées "
             "de manière neutre avec renvoi aux articles sources (le texte du règlement fait foi)."))
    _article3_table(doc, A3)
    doc.add_paragraph()

    # Article 4
    _p(doc, "Article 4 – Servitudes d’utilité publique (SUP)", bold=True)
    _kv(doc, "Liste des SUP identifiées", A4["liste_SUP"])
    _p(doc, "Avertissement : seuls les actes de servitudes publiés (et leurs annexes cartographiques) font foi.")
    doc.add_paragraph()

    # Article 5
    _p(doc, "Article 5 – Risques, protections environnementales et informations complémentaires", bold=True)
    _kv(doc, "PPR (inondation, mouvements, etc.)", A5["ppr_detail"])
    _kv(doc, "Retrait-gonflement des argiles (RGA)", A5["rga_detail"])
    _kv(doc, "Zonage sismique", A5["sismique_detail"])
    _kv(doc, "Expositions environnementales (ex. ZNIEFF, Natura 2000)", A5["env_detail"])
    _kv(doc, "Autres informations utiles", A5["autres_infos"])
    doc.add_paragraph()

    # Article 6
    _p(doc, "Article 6 – Équipements publics et réseaux", bold=True)
    _kv(doc, "État des équipements existants/prévus (AEP, assainissement, électricité, communications)", A6["etat_equipements"])
    doc.add_paragraph()

    # Article 7
    _p(doc, "Article 7 – Taxes et participations d’urbanisme", bold=True)
    _kv(doc, "Taxe d’aménagement (TA)", A7["detail_TA"])
    _kv(doc, "Redevance d’archéologie préventive (RAP)", A7["detail_RAP"])
    _kv(doc, "Autres participations", A7["autres_participations"])
    doc.add_paragraph()

    # Article 8
    _p(doc, "Article 8 – Droit de préemption", bold=True)
    _kv(doc, "Périmètre soumis au DPU", A8["dpu_detail"])
    _p(doc, "Toute mutation à titre onéreux dans ce périmètre doit faire l’objet d’une DIA.")
    doc.add_paragraph()

    # Article 9
    _p(doc, "Article 9 – Validité et portée du certificat", bold=True)
    _p(doc, ("Le présent certificat est valable 18 mois à compter de sa signature. Il gèle les règles "
             "d’urbanisme listées durant ce délai, sans préjudice de l’évolution des SUP, des risques "
             "naturels/technologiques, ni des règles fiscales et des équipements publics qui ne sont pas "
             "figés par un CU a."))
    doc.add_paragraph()

    # Article 10
    _p(doc, "Article 10 – Voies et délais de recours", bold=True)
    _p(doc, ("Recours gracieux dans les 2 mois à compter de la notification et/ou recours contentieux "
             "devant le tribunal administratif compétent dans le même délai."))
    doc.add_paragraph()

    # Fait à
    _p(doc, f"Fait à {F['commune']}, le {F['date_arrete']}")
    _p(doc, "Le Maire")
    _p(doc, "[Nom – Signature]")
    doc.add_paragraph()

    # Annexes
    _p(doc, "Annexes (à joindre à la notification)", bold=True)
    _p(doc, "1) Plan de localisation du terrain (polygone WGS84)")
    _p(doc, "2) Extraits thématiques (zonage PLU, PPR, EBC, SUP, DPU, etc.)")
    _p(doc, "3) Articles du règlement du PLU cités (extraits PDF)")

    doc.save(out_path)
    logger.info(f"✅ DOCX (brut) écrit: {out_path}")
    return out_path


# =========================
# 3) DOCX — LLM (seul l’article 3 “Dispositions” peut être enrichi)
# =========================
_PROMPT_A3 = """Tu écris des extraits STRICTEMENT factuels (sans rien inventer) pour les 9 thématiques PLU.
Ta seule source pour les règles est le texte du règlement fourni ci-dessous (extraits par zone du PLU).
Si une information n'est pas présente dans ce texte, réponds "N/D".

Retourne du JSON compact avec exactement ces clés:
{{
 "implantation_voies": {{"articles":"...", "disp":"..."}},
 "implantation_limites": {{"articles":"...", "disp":"..."}},
 "hauteurs": {{"articles":"...", "disp":"..."}},
 "emprise": {{"articles":"...", "disp":"..."}},
 "stationnement": {{"articles":"...", "disp":"..."}},
 "espaces_verts": {{"articles":"...", "disp":"..."}},
 "acces_voirie": {{"articles":"...", "disp":"..."}},
 "aspect": {{"articles":"...", "disp":"..."}},
 "annexes": {{"articles":"...", "disp":"..."}}
}}

Contexte synthétique (CUModel — ne sert qu'à connaître les zones impactant la parcelle):
{cumodel_json}

Texte du règlement PLU (par zone):
{plu_reglement}
"""

def _safe_json_loads(s: str) -> dict:
    try:
        return json.loads(s)
    except Exception:
        return {}

def generate_docx_llm(cumodel: Dict[str, Any], out_path: str) -> str:
    # On part du brut, mais on tente de remplir l'article 3 via LLM.
    a3 = cumodel.get("article3", {}).copy()

    # Règlement PLU brut par zone (si dispo)
    reg_txt = ""
    if _PLU_REG_OK:
        try:
            # Log des zones PLU trouvées dans le cumodel
            plu_zones = (cumodel.get("plu") or {}).get("zones") or []
            logger.info(f"🔍 Zones PLU trouvées dans le cumodel: {len(plu_zones)} zones")
            
            for i, zone in enumerate(plu_zones):
                libelle = zone.get("libelle", "N/A")
                typezone = zone.get("typezone", "N/A")
                pct = zone.get("pct", "N/A")
                logger.info(f"  Zone {i+1}: libelle='{libelle}', typezone='{typezone}', pct={pct}")
            
            zones_to_text = get_regulation_for_cumodel_zones(cumodel)
            logger.info(f"📋 Règlements récupérés pour {len(zones_to_text)} zones")
            
            # mapping % par code (optionnel)
            pct_by_zone = {}
            for z in plu_zones:
                # essaie de rattacher le % au code le plus précis
                from plu_regulation_context import infer_zone_codes
                codes = infer_zone_codes(z.get("libelle"), z.get("typezone"))
                logger.info(f"  Codes inférés pour zone '{z.get('libelle')}': {codes}")
                for c in codes:
                    pct_by_zone[c] = z.get("pct")
                    logger.info(f"    Code '{c}' -> pct {z.get('pct')}")
            
            logger.info(f"📊 Mapping final pct_by_zone: {pct_by_zone}")
            reg_txt = join_regulations_for_docx(zones_to_text, pct_by_zone)
            logger.info(f"📝 Texte règlement final généré ({len(reg_txt)} caractères)")
            
        except Exception as e:
            logger.warning(f"Erreur lors de la récupération du règlement PLU: {e}")
            reg_txt = ""

    if _LLM_OK:
        prompt = _PROMPT_A3.format(
            cumodel_json=json.dumps(cumodel, ensure_ascii=False),
            plu_reglement=reg_txt or "(aucun texte de règlement fourni)"
        )
        res = call_gpt5_text(prompt, reasoning_effort="low", verbosity="low")
        if res.get("success"):
            txt = (res.get("response") or "").strip()
            parsed = _safe_json_loads(txt)
            def fill(k_key, art_key, disp_key):
                block = parsed.get(k_key) or {}
                a3[art_key] = block.get("articles") or a3.get(art_key) or ""
                a3[disp_key] = block.get("disp") or a3.get(disp_key) or ""
            fill("implantation_voies", "articles_implantation_voies", "disp_implantation_voies")
            fill("implantation_limites", "articles_implantation_limites", "disp_implantation_limites")
            fill("hauteurs", "articles_hauteurs", "disp_hauteurs")
            fill("emprise", "articles_emprise", "disp_emprise")
            fill("stationnement", "articles_stationnement", "disp_stationnement")
            fill("espaces_verts", "articles_espaces_verts", "disp_espaces_verts")
            fill("acces_voirie", "articles_acces_voirie", "disp_acces_voirie")
            fill("aspect", "articles_aspect", "disp_aspect")
            fill("annexes", "articles_annexes", "disp_annexes")
        else:
            logger.warning(f"LLM indisponible pour Article 3: {res.get('error')}")

    # Écrit le DOCX avec le même gabarit mais Article 3 enrichi si possible
    cm2 = dict(cumodel)
    cm2["article3"] = a3
    return generate_docx_brut(cm2, out_path)
