#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
cua_builder.py — Génère le CUA (DOCX) final — Modèle v3.

Structure :
- En-tête 1ʳᵉ page (cua_header.py) : logo, titres, tableau récap 1/2 page + QR
- Article 1 : Objet
- Article 2 : Identification et localisation du terrain (tableau)
- Article 3 : Dispositions d’urbanisme (Zonage)  ← ajoute "Zonage : N"
- Article 4 : Servitudes d’utilité publique (SUP)
- Article 5 : Risques & protections environnementales (PPR, RGA, sismique, expositions)
- Article 6 : Équipements publics et réseaux
- Article 7 : Taxes & participations
- Article 8 : Droit de préemption
- Annexes & informations légales
- Annexe PLU (optionnelle, nettoyée via GPT-5 Nano)
- Pied de page : numéro de dossier sur toutes les pages
- Signature : dernière page, logo inséré dans le corps (pas en footer)
"""

import os, argparse
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ---------- LLM (optionnel) — pour Observations automatiques si besoin ----------
_HAS_LLM = False
try:
    import path_setup  # Configure le Python path
    from UTILS.llm_utils import call_gpt5_text, call_gpt5_nano  # noqa: F401
    _HAS_LLM = True
except Exception:
    _HAS_LLM = False

# ---------- PLU annex (optionnel) ----------
_HAS_PLU = False
try:
    from .fetch_plu_regulation import (  # noqa: F401
        canonicalize_zone,
        fetch_plu_regulations_for_zones,
        join_regulations_for_docx,
    )
    _HAS_PLU = True
except Exception:
    _HAS_PLU = False

# ---------- En-tête / intro (module externe, inchangé) ----------
from .cua_header import (  # type: ignore
    setup_first_page_header,
    add_centered_titles,
    render_intro_block_with_qr,
    add_mayor_section_with_vu,
)

# ---------- Utils (petites fonctions externalisées) ----------
from .cua_builder_utils import (
    read_json, date_fr, safe, join_addr,
    parcels_label_from_cerfa, terrain_addr_from_cerfa, demandeur_block, format_footer_numero,
    extract_zones_and_pct, extract_sup_list, build_ppr_detail, build_rga_detail, build_sismique_detail,
    build_env_detail, build_other_infos,
    build_ppr_struct, pct_fr, build_env_struct, group_parcels_by_value_pct
)



# --------------------- DOCX helpers (spécifiques au builder) ---------------------

ARTICLE_SPACE_AFTER_PT = 14

def _setup_doc() -> Document:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Urbanist"
    # compat Word/Office
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Urbanist')
    st.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2); s.right_margin = Cm(2)
        s.different_first_page_header_footer = True
    return doc

def _set_footer_num(doc: Document, text: str):
    for s in doc.sections:
        p = s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.clear(); p.add_run(text).italic = True

def add_article_title(doc: Document, title_text: str):
    p = doc.add_paragraph()
    run = p.add_run(title_text.upper()); run.bold = True; run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p

def add_paragraph(doc: Document, text: str, *, bold: bool=False, italic: bool=False, center=False):
    p = doc.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.italic = italic
    p.paragraph_format.space_after = Pt(ARTICLE_SPACE_AFTER_PT//2)
    return p

def add_legal_paragraph(doc: Document, text: str, *, italic: bool=False):
    """Ajoute un paragraphe avec une police plus petite pour les textes légaux"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(8)  # Plus petit que la taille normale (11pt)
    p.paragraph_format.space_after = Pt(ARTICLE_SPACE_AFTER_PT//3)
    return p

def add_kv_table(doc: Document, rows: List[Tuple[str,str]]):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid"
    try:
        t.columns[0].width = Cm(5.0)
    except Exception:
        pass
    for i, (k,v) in enumerate(rows):
        p0 = t.cell(i,0).paragraphs[0]; p0.add_run(k)
        p1 = t.cell(i,1).paragraphs[0]; r1 = p1.add_run(v or "—"); r1.bold = True
        for par in (p0, p1):
            par.space_after = Pt(2); par.space_before = Pt(2)
    doc.add_paragraph("")  # un peu d'air


# --------------------- Construction du CUA ---------------------

def build_cua_docx(
    cerfa_json: Dict[str, Any],
    intersections_json: Dict[str, Any],
    output_docx: str,
    *,
    logo_first_page: Optional[str] = None,
    signature_logo: Optional[str] = None,
    include_plu_annex: bool = True,
    # meta PLU (si tu veux surcharger)
    plu_nom: str = "PLU en vigueur",
    plu_date_appro: str = "13/02/2017",
) -> None:

    data = cerfa_json.get("data") or {}
    commune = (data.get("commune_nom") or "").upper()

    # Champs dynamiques
    date_dep = date_fr(data.get("date_depot"))
    who, domicile = demandeur_block(cerfa_json)
    terrain = terrain_addr_from_cerfa(cerfa_json)
    parcelles = parcels_label_from_cerfa(cerfa_json)
    surface = str(data.get("superficie_totale_m2") or "").strip()
    num_cu = data.get("numero_cu") or ""
    footer_num = format_footer_numero(cerfa_json)

    # Build
    doc = _setup_doc()

    # 1) En-tête / titres / tableau récap + QR
    setup_first_page_header(doc.sections[0], commune, logo_first_page)
    add_centered_titles(doc)
    render_intro_block_with_qr(doc, cerfa_json, qr_text=num_cu, half_page=True, 
                              qr_logo_path="/Users/benjaminbenoit/Downloads/k.png")

    # Pied de page
    _set_footer_num(doc, footer_num)

    # Section "Le Maire" avec "Vu" et "CERTIFIE"
    add_mayor_section_with_vu(doc, cerfa_json, commune, plu_date_appro)

    # -------------------- Articles --------------------

    # Article 1 — Objet
    add_article_title(doc, "Article UN - Objet")
    add_paragraph(doc,
        "Les règles d’urbanisme, la liste des taxes et participations d’urbanisme ainsi que les limitations administratives au droit de propriété applicables au terrain sont mentionnées aux articles 2 et suivants du présent certificat.")
    add_paragraph(doc,
        "Conformément au quatrième alinéa de l’article L. 410-1 du code de l’urbanisme, si une demande de permis de construire, d’aménager ou de démolir ou si une déclaration préalable est déposée dans le délai de dix-huit mois à compter de la date du présent certificat d'urbanisme, les dispositions d'urbanisme, le régime des taxes et participations d'urbanisme ainsi que les limitations administratives au droit de propriété tels qu'ils existaient à cette date ne peuvent être remis en cause à l'exception des dispositions qui ont pour objet la préservation de la sécurité ou de la salubrité publique.")

    # Article 2 — Identification & localisation
    add_article_title(doc, "Article DEUX - Identification et localisation du terrain")
    add_kv_table(doc, [
        ("Commune", f"{(data.get('commune_nom') or '').title()} ({data.get('commune_insee') or ''})"),
        ("Adresse / Localisation", terrain),
        ("Références cadastrales", parcelles),
        ("Surface indicative", (surface + " m²") if surface else "—"),
        ("Document d’urbanisme opposable", f"{plu_nom} approuvé le {plu_date_appro}"),
    ])

    # Article 3 — Dispositions d'urbanisme (Zonage)
    add_article_title(doc, "Article TROIS - Dispositions d'urbanisme applicables")
    zones, pct_by = extract_zones_and_pct(intersections_json)

    # Ligne "Zonage : X, Y..."
    if zones:
        add_paragraph(doc, "Zonage : " + ", ".join(zones), bold=True)
    else:
        add_paragraph(doc, "Zonage : non déterminé", bold=True)

    # Extraire les numéros de parcelles (sans section)
    refs = (data.get("references_cadastrales") or [])
    parcelle_nums = [str(r.get("numero") or "").zfill(3) for r in refs if r.get("numero")]
    parcelles_str_nums = ", ".join(parcelle_nums) if parcelle_nums else "—"

    pz = doc.add_paragraph()
    rz = pz.add_run("Zonage du Plan Local d'Urbanisme (PLU) : "); rz.bold = True
    if parcelle_nums:
        pz.add_run(f"Parcelles {parcelles_str_nums} : ")
    pz.add_run(", ".join(zones) if zones else "non déterminé").bold = True
    pz.paragraph_format.space_after = Pt(ARTICLE_SPACE_AFTER_PT)

    # Article 4 — SUP
    add_article_title(doc, "Article QUATRE - Servitudes d’utilité publiques (SUP)")
    sup_list = extract_sup_list(intersections_json)
    if sup_list:
        for sup, nom in sup_list:
            doc.add_paragraph(f"- {sup or '—'} : {nom or '—'}")
    else:
        doc.add_paragraph("— Aucune SUP détectée dans les données fournies.")
    add_paragraph(doc, "Avertissement : seuls les actes de servitudes publiés (et leurs annexes cartographiques) font foi.", italic=True)

    # Article 5 — Risques & protections environnementales
    add_article_title(doc, "Article CINQ – Risques, protections environnementales et informations complémentaires")

    # ---- 5.1) PPR (inondation, mouvements, etc.) ----
    from .cua_builder_utils import build_ppr_struct, pct_fr
    ppr = build_ppr_struct(intersections_json)

    # Titre de partie en gras
    p = doc.add_paragraph()
    p.add_run("1) PPR (inondation, mouvements, etc.)").bold = True

    # Synthèse inter-parcelles (Zonage réglementaire)
    if ppr.get("zonage"):
        add_paragraph(doc, "Synthèse inter-parcelles", italic=True)
        pz = doc.add_paragraph()
        pz.add_run("Zonage réglementaire PPRI (codezone) :").bold = True
        for pnum in sorted(ppr["zonage"].keys(), key=lambda x: (len(x), x)):
            pairs = ppr["zonage"][pnum]
            line = " • Parcelle " + pnum + " : " + ", ".join([f"{v} ({pct_fr(p)} %)" for v, p in pairs])
            doc.add_paragraph(line)

    # Synthèse inter-parcelles (Isocotes)
    if ppr.get("isocotes"):
        doc.add_paragraph("")  # saut de ligne
        pi = doc.add_paragraph()
        pi.add_run("Isocotes (cotes de référence PPRI) :").bold = True
        for pnum in sorted(ppr["isocotes"].keys(), key=lambda x: (len(x), x)):
            pairs = ppr["isocotes"][pnum]
            line = " • Parcelle " + pnum + " : " + ", ".join([f"{v} ({pct_fr(p)} %)" for v, p in pairs])
            doc.add_paragraph(line)

    # (IMPORTANT) On NE répète PAS le détail : la synthèse ci-dessus suffit et évite le doublon.
    # Sources PPR
    srcs = ppr.get("sources") or set()
    if srcs:
        doc.add_paragraph(f"Source PPRI (documents et découpages) : " + "; ".join(sorted(srcs)))

    doc.add_paragraph("")  # grand saut de ligne entre sous-parties

    # ---- 5.2) RGA ----
    pr = doc.add_paragraph()
    pr.add_run("2) Retrait-gonflement des argiles (RGA)").bold = True
    add_paragraph(doc, build_rga_detail(intersections_json))

    doc.add_paragraph("")  # saut

    # ---- 5.3) Zonage sismique ----
    ps = doc.add_paragraph()
    ps.add_run("3) Tremblement de terre").bold = True
    add_paragraph(doc, build_sismique_detail(intersections_json))

    doc.add_paragraph("")  # saut

    # ---- 5.4) Expositions environnementales ----
    from .cua_builder_utils import build_env_struct, group_parcels_by_value_pct
    pe = doc.add_paragraph()
    pe.add_run("4) Expositions environnementales (ZNIEFF, Natura 2000, radon, nuisances, etc.)").bold = True

    envs = build_env_struct(intersections_json)

    # Synthèse inter-parcelles (regroupée)
    if envs.get("radon") or envs.get("nuisances"):
        add_paragraph(doc, "Synthèse inter-parcelles", italic=True)

    # Radon — regrouper les parcelles par (valeur, pct)
    if envs.get("radon"):
        label = envs["radon"].get("_label") or "Radon"
        pradon = doc.add_paragraph()
        pradon.add_run(f"{label} :").bold = True

        rows = group_parcels_by_value_pct({k:v for k,v in envs["radon"].items() if k != "_label"}, with_trunk=False)
        if rows:
            for (val, pct, _, plist) in rows:
                doc.add_paragraph(" • Parcelles " + ", ".join(plist) + f" : {val} ({pct_fr(pct)} %)")

    # Nuisances — regrouper par (valeur, pct, axe)
    if envs.get("nuisances"):
        doc.add_paragraph("")  # saut de ligne
        label = envs["nuisances"].get("_label") or "Nuisances sonores"
        pnuis = doc.add_paragraph()
        pnuis.add_run(f"{label} :").bold = True

        rows = group_parcels_by_value_pct({k:v for k,v in envs["nuisances"].items() if k != "_label"}, with_trunk=True)
        if rows:
            for (val, pct, trunk, plist) in rows:
                extra = f" – axe {trunk}" if trunk else ""
                doc.add_paragraph(" • Parcelles " + ", ".join(plist) + f" : {val} ({pct_fr(pct)} %){extra}")

    # Détail par parcelle (optionnel) — on peut garder si tu y tiens,
    # mais comme la synthèse est groupée, c'est souvent suffisant visuellement.
    # Ici, on supprime le Détail pour éviter toute redite.
    # Si tu veux le réactiver plus tard, on pourra l'ajouter sous condition.

    doc.add_paragraph("")  # saut

    # ---- 5.5) Autres informations utiles ----
    pau = doc.add_paragraph()
    pau.add_run("5) Autres informations utiles").bold = True
    add_paragraph(doc, build_other_infos(intersections_json))

    # Article 6 — Équipements & réseaux
    add_article_title(doc, "Article SIX – Équipements publics et réseaux")
    add_paragraph(doc, "État des équipements existants/prévus (AEP, assainissement, électricité, communications) : Non renseigné.")

    # Article 7 — Taxes & participations
    add_article_title(doc, "Article SEPT – Taxes et participations d’urbanisme")
    doc.add_paragraph(
        "Les taxes suivantes pourront être exigées à compter de l'obtention d'un permis ou d'une décision de non opposition à une déclaration préalable.\n\n"
        "Taxe d’Aménagement :\n"
        "Part communale : Taux : 5% \n"
        "Part départementale : Taux : 2,5 % \n"
        "Redevance d’Archéologie Préventive : Taux : 0,68 %"
    ).paragraph_format.space_after = Pt(ARTICLE_SPACE_AFTER_PT)
    doc.add_paragraph(
        "Les participations ci-dessous pourront être exigées à l'occasion d'un permis de construire ou d'une décision de non opposition à une déclaration préalable. "
        "Si tel est le cas elles seront mentionnées dans l'arrêté de permis ou dans un arrêté pris dans les deux mois suivant la date du permis tacite ou de la décision de non opposition à une déclaration préalable.\n\n"
        "Participations susceptibles d’être exigés à l’occasion de l’opération :\n"
        "- contribution aux dépenses de réalisation des équipements publics.\n"
        "- financement de branchements des équipements propres (article L332-15 du CU)."
    ).paragraph_format.space_after = Pt(ARTICLE_SPACE_AFTER_PT)

    # Article 8 — Droit de préemption
    add_article_title(doc, "Article HUIT – Droit de préemption")
    # Ici, on pourrait croiser un type 'preemption' si tu veux une phrase dynamique :
    # preempt = extract_preemption(intersections_json)  # à implémenter si besoin
    # if preempt: ...
    doc.add_paragraph("Le terrain n'est pas situé dans une zone de droit de préemption.")
    doc.add_paragraph("Aucune DIA (Déclaration d'Intention d'Aliéner) au titre du DPU n'est requise.")


    # Annexes & informations à lire attentivement
    add_article_title(doc, "ANNEXES")
    doc.add_paragraph("1) Plan de localisation du terrain")
    doc.add_paragraph("2) Extraits thématiques")
    doc.add_paragraph("3) Articles du règlement du PLU")

    add_article_title(doc, "INFORMATIONS À LIRE ATTENTIVEMENT")
    add_legal_paragraph(doc,
        "Le (ou les) demandeur(s) peut contester la légalité de la décision dans les deux mois qui suivent la date de sa notification. "
        "Durée de validité : 18 mois, prorogeable par périodes d'un an sous conditions (art. R. 410-17-1). "
        "Le certificat d'urbanisme est un acte d'information et ne vaut pas autorisation. "
        "En cas de dépôt d'une autorisation dans le délai de validité, les nouvelles dispositions ne pourront pas être opposées, sauf exceptions liées à la sécurité ou à la salubrité publique.",
        italic=True
    )
    add_legal_paragraph(doc,
        "Le (ou les) demandeur(s) peut contester la légalité de la décision dans les deux mois qui suivent la date de sa notification. A cet effet il peut saisir le tribunal administratif territorialement compétent d'un recours contentieux. "
        "Durée de validité : Le certificat d'urbanisme a une durée de validité de 18 mois. Il peut être prorogé par périodes d'une année si les prescriptions d'urbanisme, les servitudes d'urbanisme de tous ordres et le régime des taxes et participations n'ont pas évolué. Vous pouvez présenter une demande de prorogation en adressant une demande sur papier libre, accompagnée du certificat pour lequel vous demandez la prorogation au moins deux mois avant l'expiration du délai de validité. "
        "A défaut de notification d'une décision expresse portant prorogation du certificat d'urbanisme dans le délai de deux mois suivant la réception en mairie de la demande, le silence gardé par l'autorité compétente vaut prorogation du certificat d'urbanisme. La prorogation prend effet au terme de la validité de la décision initiale (Art. R. 410-17-1). "
        "Effets du certificat d'urbanisme : le certificat d'urbanisme est un acte administratif d'information, qui constate le droit applicable en mentionnant les possibilités d'utilisation de votre terrain et les différentes contraintes qui peuvent l'affecter. Il n'a pas valeur d'autorisation pour la réalisation des travaux ou d'une opération projetée. "
        "Le certificat d'urbanisme crée aussi des droits à votre égard. Si vous déposez une demande d'autorisation (par exemple une demande de permis de construire) dans le délai de validité du certificat, les nouvelles dispositions d'urbanisme ou un nouveau régime de taxes ne pourront pas vous être opposées, sauf exceptions relatives à la préservation de la sécurité ou de la salubrité publique.",
        italic=True
    )

    add_legal_paragraph(doc,
        "QR Code : permet d'accéder à une Carte interactive des règles applicables (zonage, SUP, risques, prescriptions, obligations, informations). "
        "Affichage informatif ; en cas de divergence, les pièces écrites et le règlement en vigueur font foi. "
        "Solution proposée par KERELIA (RCS Bordeaux 944 763 275).",
        italic=True
    )

    # Annexe PLU (facultative)
    if include_plu_annex and _HAS_PLU and zones:
        doc.add_page_break()
        add_paragraph(doc, "ANNEXE — Règlement PLU", bold=True, center=True)
        zones_text = fetch_plu_regulations_for_zones(
            zones,
            table="plu_regulations_clean",  # Nom de la nouvelle table
            debug=False,
        )
        if zones_text:
            bloc = join_regulations_for_docx(
                zones_text,
                pct_by_zone={k: pct_by.get(k, 0.0) for k in zones_text.keys()}
            )
            for chunk in (bloc or "").split("\n\n"):
                chunk = chunk.strip()
                if chunk:
                    doc.add_paragraph(chunk)
        else:
            doc.add_paragraph("Aucune réglementation PLU retrouvée en base pour les zones détectées.")

    # Signature (dernière page, logo dans le corps)
    from datetime import datetime as _dt
    doc.add_page_break()
    add_paragraph(doc, f"Fait à {commune or 'LATRESNE'},", bold=True)
    add_paragraph(doc, f"Le {_dt.now().strftime('%d/%m/%Y')}")
    add_paragraph(doc, "Le Maire,")
    if signature_logo and os.path.exists(signature_logo):
        pimg = doc.add_paragraph(); pimg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            pimg.add_run().add_picture(signature_logo, width=Cm(3.0))
        except Exception:
            pass

    doc.save(output_docx)


# --------------------- CLI ---------------------

def main():
    ap = argparse.ArgumentParser(description="CUA DOCX (modèle v3).")
    ap.add_argument("--cerfa-json", required=True)
    ap.add_argument("--intersections-json", required=True)
    ap.add_argument("--output", default="CUA_final.docx")
    ap.add_argument("--logo-first-page", default="/Volumes/T7/Travaux_Freelance/KERELIA/CUAs/DATA_FLAVIO/INTERSECTION_CALCULS/INTERSECTION_APP/commune_de_latresne.png")
    ap.add_argument("--signature-logo", default="/Users/benjaminbenoit/Downloads/k.png")
    ap.add_argument("--no-plu-annex", action="store_true")
    # éventuels overrides d'intitulé PLU
    ap.add_argument("--plu-nom", default="PLU en vigueur")
    ap.add_argument("--plu-date-appro", default="13/02/2017")
    args = ap.parse_args()

    cerfa = read_json(args.cerfa_json)
    inters = read_json(args.intersections_json)

    build_cua_docx(
        cerfa, inters, args.output,
        logo_first_page=args.logo_first_page,
        signature_logo=args.signature_logo,
        include_plu_annex=(not args.no_plu_annex),
        plu_nom=args.plu_nom,
        plu_date_appro=args.plu_date_appro,
    )
    print(f"✅ CUA généré : {args.output}")


if __name__ == "__main__":
    main()
