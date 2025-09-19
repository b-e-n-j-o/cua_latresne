# -*- coding: utf-8 -*-
"""
cua_header.py — En-tête CUA (1ʳᵉ page)
- Logo haut-gauche + "MAIRIE DE …" centré (header 1ʳᵉ page uniquement)
- Titres centrés
- Tableau récap à gauche (moitié de page) + QR code à droite (hors tableau)
"""

import os, io, datetime
from typing import Any, Tuple
from docx.document import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- Helpers données CERFA ----------
def _date_fr(iso: str | None) -> str:
    if not iso:
        return ""
    try:
        d = datetime.date.fromisoformat(iso)
        return d.strftime("%d/%m/%Y")
    except Exception:
        return iso

def _safe(x, default=""):
    return default if x in (None, "", []) else x

def _join_addr(ad: dict) -> str:
    if not ad: return ""
    parts = []
    if ad.get("numero"): parts.append(str(ad["numero"]).strip())
    if ad.get("voie"): parts.append(str(ad["voie"]).strip())
    if ad.get("lieu_dit"): parts.append(str(ad["lieu_dit"]).strip())
    line1 = " ".join(parts).strip()
    line2 = " ".join([_safe(ad.get("code_postal")), _safe(ad.get("ville"))]).strip()
    return (line1 + (", " + line2 if line2 else "")).strip()

def _demandeur_block(cerfa: dict) -> Tuple[str, str]:
    d = (cerfa.get("data") or {}).get("demandeur") or {}
    who = (d.get("denomination") or " ".join([_safe(d.get("prenom")), _safe(d.get("nom"))]).strip()).strip()
    siret = _safe(d.get("siret"))
    who_fmt = (who.upper() + (f" (SIRET {siret})" if siret else ""))
    domicile = _join_addr(((cerfa.get("data") or {}).get("coord_demandeur") or {}).get("adresse") or {})
    return who_fmt, domicile

def _terrain_addr(cerfa: dict) -> str:
    return _join_addr(((cerfa.get("data") or {}).get("adresse_terrain") or {}))

# ---------- Helpers mise en page ----------
def _emu_to_cm(v: int) -> float:
    # 1 cm = 360000 EMU
    return float(v) / 360000.0

def _content_width_cm(section) -> float:
    return _emu_to_cm(section.page_width - section.left_margin - section.right_margin)

# ---------- Header 1ʳᵉ page ----------
def setup_first_page_header(section, commune_name: str, logo_path: str | None):
    """Logo à gauche + 'MAIRIE DE …' centré, uniquement sur la 1ʳᵉ page."""
    section.different_first_page_header_footer = True
    hdr = section.first_page_header

    # ligne 1 : logo à gauche
    p1 = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path and os.path.exists(logo_path):
        try:
            run = p1.add_run()
            run.add_picture(logo_path, width=Cm(3.0))
        except Exception:
            pass

    # ligne 2 : texte centré
    p2 = hdr.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(f"MAIRIE DE {commune_name.upper()}")
    r.bold = True
    r.font.size = Pt(14)

# ---------- Titres centrés ----------
def add_centered_titles(doc: Document):
    t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run("CERTIFICAT D’URBANISME"); r1.bold = True; r1.font.size = Pt(20)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.add_run("Délivré par le maire au nom de la commune").italic = True

# ---------- QR ----------
def _make_qr_png_bytes(text: str, box_size: int = 12, border: int = 2, logo_path: str | None = None) -> bytes:
    """
    Génère un QR PNG en mémoire avec logo optionnel au centre.
    Requiert qrcode + pillow. Si indisponibles, renvoie un placeholder blanc.
    """
    try:
        import qrcode
        from PIL import Image, ImageDraw
        
        # Créer le QR code avec niveau de correction élevé pour tolérer le logo
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_H,  # Niveau élevé (30% d'erreur tolérée)
            box_size=box_size,
            border=border
        )
        qr.add_data(text or "")
        qr.make(fit=True)
        
        # Générer l'image QR de base
        qr_img = qr.make_image(fill_color="black", back_color="white").convert('RGB')
        
        # Ajouter le logo au centre si fourni
        if logo_path and os.path.exists(logo_path):
            try:
                # Charger et redimensionner le logo
                logo = Image.open(logo_path).convert("RGBA")
                
                # Calculer la taille du logo (plus grand : environ 25% de la taille du QR)
                qr_width, qr_height = qr_img.size
                logo_size = min(qr_width, qr_height) // 4  # Plus grand que avant (1/4 au lieu de 1/6)
                
                # Forcer le logo à être carré en redimensionnant avec déformation si nécessaire
                logo_square = logo.resize((logo_size, logo_size), Image.Resampling.LANCZOS)
                
                # Calculer la position centrale
                pos_x = (qr_width - logo_size) // 2
                pos_y = (qr_height - logo_size) // 2
                
                # Coller le logo carré directement au centre du QR code (avec transparence si RGBA)
                if logo_square.mode == 'RGBA':
                    qr_img.paste(logo_square, (pos_x, pos_y), logo_square)
                else:
                    qr_img.paste(logo_square, (pos_x, pos_y))
                
            except Exception as e:
                # Si erreur avec le logo, continuer sans logo
                pass
        
        # Sauvegarder en PNG
        buf = io.BytesIO()
        qr_img.save(buf, format="PNG")
        return buf.getvalue()
        
    except Exception:
        # Fallback : placeholder blanc
        from PIL import Image
        img = Image.new("RGB", (512, 512), "white")
        buf = io.BytesIO(); img.save(buf, format="PNG")
        return buf.getvalue()

# ---------- Bloc récap (tableau) + QR à droite ----------
def render_intro_block_with_qr(
    doc: Document,
    cerfa: dict,
    qr_text: str,
    *,
    half_page: bool = True,
    left_width_cm: float | None = None,
    right_width_cm: float | None = None,
    gap_cm: float = 0.0,
    qr_logo_path: str | None = None
) -> None:
    """
    Tableau récap à gauche (moitié de la largeur utile), QR centré à droite (hors tableau).
    """
    data = cerfa.get("data") or {}
    date_dep = _date_fr(data.get("date_depot"))
    who, domicile = _demandeur_block(cerfa)
    terrain = _terrain_addr(cerfa)
    num_cu  = data.get("numero_cu") or ""

    # Largeurs 50/50 si non précisées
    if half_page or left_width_cm is None or right_width_cm is None:
        cw = _content_width_cm(doc.sections[0])
        lcm = (cw - gap_cm) / 2.0
        rcm = (cw - gap_cm) / 2.0
    else:
        lcm, rcm = float(left_width_cm), float(right_width_cm)

    # Table "layout" 1×2 pour séparer gauche/droite
    layout = doc.add_table(rows=1, cols=2)
    layout.autofit = False
    layout.columns[0].width = Cm(lcm)
    layout.columns[1].width = Cm(rcm)

    # ---- 1) à gauche : tableau récap (labels/valeurs)
    left = layout.cell(0, 0)
    recap = left.add_table(rows=5, cols=2)
    recap.style = "Light Grid"
    recap.autofit = False
    recap.columns[0].width = Cm(5.0)
    recap.columns[1].width = Cm(max(2.0, lcm - 5.2))

    rows = [
        ("Demande déposée le :", date_dep),
        ("Par :", who),
        ("Demeurant à :", domicile),
        ("Sur un terrain sis :", terrain),
        ("Numéro du CU :", num_cu),
    ]
    for i, (label, value) in enumerate(rows):
        c0 = recap.cell(i, 0).paragraphs[0]; c0.add_run(label)
        p1 = recap.cell(i, 1).paragraphs[0]; r1 = p1.add_run(value or "—")
        r1.bold = True

    # ---- 2) à droite : QR centré + légende
    right = layout.cell(0, 1)
    pqr = right.paragraphs[0]; pqr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Générer le QR avec logo au centre si fourni
    qr_png = _make_qr_png_bytes(qr_text, logo_path=qr_logo_path)
    stream = io.BytesIO(qr_png)
    qr_width_cm = max(4.5, min(rcm * 0.85, 8.0))
    run = pqr.add_run(); run.add_picture(stream, width=Cm(qr_width_cm))
    cap = right.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_mayor_section_with_vu(doc: Document, cerfa: dict, commune: str, plu_date_appro: str = "13/02/2017"):
    """
    Ajoute la section 'Le Maire' avec tous les 'Vu' et 'CERTIFIE' 
    en utilisant les données dynamiques du CERFA.
    """
    # Extraire les données du CERFA
    data = cerfa.get("data") or {}
    date_dep = _date_fr(data.get("date_depot"))
    who, _ = _demandeur_block(cerfa)
    terrain = _terrain_addr(cerfa)
    parcelles = ""
    refs = (data.get("references_cadastrales") or [])
    if refs:
        parcelle_list = [f'{(r.get("section") or "").upper()} {str(r.get("numero") or "").zfill(4)}' for r in refs]
        parcelles = ", ".join([p for p in parcelle_list if p.strip()])
    num_cu = data.get("numero_cu") or ""
    
    # LE MAIRE (centré, majuscules)
    p_maire = doc.add_paragraph()
    p_maire.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_maire = p_maire.add_run("LE MAIRE")
    r_maire.bold = True
    r_maire.font.size = Pt(14)
    
    doc.add_paragraph("")  # espace
    
    # Bloc "Vu la demande..." (dynamique)
    p_vu_demande = doc.add_paragraph()
    p_vu_demande.add_run("Vu la demande d'un certificat d'urbanisme indiquant, en application de l'article L.410-1 a) du code de l'urbanisme, "
                        "les dispositions d'urbanisme, les limitations administratives au droit de propriété et la liste des taxes et participations d'urbanisme "
                        f"applicables à un terrain situé à ")
    r_terrain = p_vu_demande.add_run(terrain or "—")
    r_terrain.bold = True
    p_vu_demande.add_run(f" (cadastré {parcelles}), présentée le ")
    r_date = p_vu_demande.add_run(date_dep or "—")
    r_date.bold = True
    p_vu_demande.add_run(" par ")
    r_who = p_vu_demande.add_run(who or "—")
    r_who.bold = True
    p_vu_demande.add_run(f", et enregistrée par la mairie de ")
    r_commune = p_vu_demande.add_run(commune.upper())
    r_commune.bold = True
    p_vu_demande.add_run(" sous le numéro ")
    r_num = p_vu_demande.add_run(num_cu.replace("-", ""))
    r_num.bold = True
    p_vu_demande.add_run(" ;")
    
    # Autres "Vu" (statiques)
    vu_texts = [
        "Vu le Code de l'Urbanisme et notamment ses articles L.410-1, R.410-1 et suivants ;",
        f"Vu le Plan Local d'urbanisme approuvé en date du {plu_date_appro} ;",
        "Vu le Plan de Prévention du risque naturel d'inondation (PPRNI) de l'Agglomération bordelaise commune de Latresne, approuvé par arrêté préfectoral du 23 février 2022 ;",
        "Vu la délibération du conseil municipal du 1er février 2024 instaurant l'obligation de déclaration préalable lors de divisions foncières situées dans les zones naturelles et les zones agricoles sur le territoire de la commune ;",
        "Vu la délibération du conseil municipal du 1er février 2024 instaurant l'obligation de déclaration préalable lors de division du foncier bâti sur l'ensemble du territoire de la commune ;",
        "Vu la délibération du conseil municipal du 1er février 2024 instaurant l'autorisation préalable de travaux conduisant à la création de locaux à usage d'habitation dite « permis de diviser » sur l'ensemble du territoire de la commune ;"
    ]
    
    for vu_text in vu_texts:
        doc.add_paragraph(vu_text)
    
    doc.add_paragraph("")  # espace
    
    # CERTIFIE (centré, majuscules)
    p_certifie = doc.add_paragraph()
    p_certifie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_certifie = p_certifie.add_run("CERTIFIE :")
    r_certifie.bold = True
    r_certifie.font.size = Pt(14)


__all__ = [
    "setup_first_page_header",
    "add_centered_titles",
    "render_intro_block_with_qr",
    "add_mayor_section_with_vu",
]
